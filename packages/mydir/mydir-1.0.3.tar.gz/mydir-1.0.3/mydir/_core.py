from __future__ import annotations

import argparse
import fnmatch
import json
import logging
import os
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Final, Iterable, List, Set

__all__ = ["mydir", "build_cli"]

_LOG = logging.getLogger("mydir")
DEFAULT_IGNORES: Final[Set[str]] = {
    ".git",
    "__pycache__",
    ".DS_Store",
    ".idea",
    ".vscode",
    "*.pyc",
    "*.pyo",
    "*.swp",
    "node_modules",
    ".eggs",
    "dist",
    "build",
    ".gitignore",
    ".env",
    "venv",
    '.pytest_cache',
    '.ipynb_checkpoints',
}

PDF_DEFAULT_PAGE_SIZE: Final = "A4"
PDF_DEFAULT_DPI: Final = 96
PDF_DEFAULT_SCALE: Final = 1.0


# ---------------------------------------------------------------------------#
# Helper functions
# ---------------------------------------------------------------------------#
def _human_size(num: int) -> str:
    for unit in ("B", "KB", "MB", "GB", "TB"):
        if num < 1024:
            return f"{num:.1f} {unit}"
        num /= 1024
    return f"{num:.1f} PB"


def _should_ignore(name: str, patterns: Set[str]) -> bool:
    """True if *name* matches any glob in *patterns* (case‑insensitive)."""
    lower = name.lower()
    return any(fnmatch.fnmatchcase(lower, pat.lower()) for pat in patterns)


def _sort_entries(
    dir_path: Path,
    entries: Iterable[Path],
    folders_first: bool,
    reverse: bool,
) -> List[Path]:
    key = (
        (lambda p: (not p.is_dir(), p.name.lower()))
        if folders_first
        else (lambda p: (p.is_dir(), p.name.lower()))
    )
    return sorted(entries, key=key, reverse=reverse)


# ---------------------------------------------------------------------------#
# Dataclass
# ---------------------------------------------------------------------------#
@dataclass
class mydir:
    """Represent a directory tree & provide export helpers."""

    root: Path = field(default_factory=lambda: Path.cwd())
    ignores: Set[str] = field(default_factory=lambda: set(DEFAULT_IGNORES))
    folders_first: bool = True
    reverse: bool = False
    show_sizes: bool = False
    _max_depth: int | None = None
    _lines: list[str] = field(default_factory=list, init=False, repr=False)

    # ---------------------------  fluent helpers  -------------------------- #
    def ignore(self, *patterns: str) -> "mydir":
        self.ignores.update(patterns)
        return self

    def files_first(self) -> "mydir":
        self.folders_first = False
        return self

    def max_depth(self, depth: int | None) -> "mydir":
        self._max_depth = depth if depth is None or depth >= 0 else None
        return self

    def show_sizes_on(self, enable: bool = True) -> "mydir":
        self.show_sizes = enable
        return self

    # ---------------------------  core builders  --------------------------- #
    def _format_entry(self, p: Path) -> str:
        if self.show_sizes and p.is_file():
            try:
                size = _human_size(p.stat().st_size)
                return f"{p.name} [{size}]"
            except OSError:
                pass
        return p.name

    def _walk(self, current: Path, prefix: str = "", depth: int = 0) -> None:
        if self._max_depth is not None and depth > self._max_depth:
            return
        try:
            items = [
                child
                for child in current.iterdir()
                if not _should_ignore(child.name, self.ignores)
            ]
        except PermissionError:
            self._lines.append(prefix + "└── [Permission denied]")
            return

        entries = _sort_entries(current, items, self.folders_first, self.reverse)
        total = len(entries)

        for i, entry in enumerate(entries):
            connector = "└── " if i == total - 1 else "├── "
            self._lines.append(prefix + connector + self._format_entry(entry))
            if entry.is_dir() and not entry.is_symlink():
                new_prefix = prefix + ("    " if i == total - 1 else "│   ")
                self._walk(entry, new_prefix, depth + 1)

    # ---------------------------  public API  ------------------------------ #
    def build_tree_str(self) -> str:
        self._lines.clear()
        self._lines.append(f"{self.root.resolve().name}/")
        self._walk(self.root)
        return "\n".join(self._lines)

    def to_json(self) -> str:
        return json.dumps(self._lines or self.build_tree_str().splitlines(), indent=2)

    # ---------------  console & persistence helpers  ----------------------- #
    def print(self) -> None:
        try:
            from rich.console import Console

            Console().print(self.build_tree_str())
        except ModuleNotFoundError:
            print(self.build_tree_str())

    def save(self, output: Path | str) -> None:
        out = Path(output)
        text = self.build_tree_str()
        if out.suffix.lower() == ".md":
            text = f"```\n{text}\n```"
        out.write_text(text, encoding="utf-8")

    # ---------------  PDF & DOCX bundling (unchanged logic)  --------------- #
    def export_code_pdf(
        self,
        pdf_path: Path | str,
        *,
        page_size: str = PDF_DEFAULT_PAGE_SIZE,
        dpi: int = PDF_DEFAULT_DPI,
        scale: float = PDF_DEFAULT_SCALE,
        theme: str = "default",
    ) -> None:
        try:
            from pygments import highlight
            from pygments.formatters import HtmlFormatter
            from pygments.lexers import guess_lexer_for_filename
            from weasyprint import CSS, HTML
        except ModuleNotFoundError as exc:
            raise ImportError(
                "PDF export requires extras: pip install mydir[pdf]"
            ) from exc

        formatter = HtmlFormatter(full=False, style=theme, linenos=True)
        css = formatter.get_style_defs(".highlight")
        html_parts: list[str] = []

        for file in sorted(self.root.rglob("*")):
            if file.is_file() and not _should_ignore(file.name, self.ignores):
                try:
                    text = file.read_text("utf-8", errors="ignore")
                    lexer = guess_lexer_for_filename(file.name, text)
                except Exception:
                    continue
                highlighted = highlight(text, lexer, formatter)
                html_parts.append(f"<h2>{file.relative_to(self.root)}</h2>{highlighted}")

        HTML(string=f"<style>{css}</style>{''.join(html_parts)}", base_url=str(self.root)).write_pdf(
            str(pdf_path),
            stylesheets=[
                CSS(
                    string=f"@page {{ size:{page_size}; margin:1cm }}"
                )
            ],
            presentational_hints=True,
            zoom=scale,
            dpi=dpi,
        )

    def export_code_docx(self, docx_path: Path | str) -> None:
        try:
            from bs4 import BeautifulSoup
            from docx import Document
            from docx.shared import Pt
            from pygments import highlight
            from pygments.formatters import HtmlFormatter
            from pygments.lexers import guess_lexer_for_filename
        except ModuleNotFoundError as exc:
            raise ImportError("DOCX export requires extras: pip install mydir[doc]") from exc

        doc = Document()
        doc.styles["Normal"].font.name = "Courier New"
        doc.styles["Normal"].font.size = Pt(8)
        formatter = HtmlFormatter(noclasses=True, style="default")

        for file in sorted(self.root.rglob("*")):
            if file.is_file() and not _should_ignore(file.name, self.ignores):
                try:
                    text = file.read_text("utf-8", errors="ignore")
                    lexer = guess_lexer_for_filename(file.name, text)
                except Exception:
                    continue
                highlighted = highlight(text, lexer, formatter)
                doc.add_heading(str(file.relative_to(self.root)), level=2)
                soup = BeautifulSoup(highlighted, "html.parser")
                for ln in soup.get_text().splitlines():
                    doc.add_paragraph(ln)
        doc.save(str(docx_path))


# ---------------------------------------------------------------------------#
# CLI
# ---------------------------------------------------------------------------#
def _add_common_tree_args(p: argparse.ArgumentParser) -> None:
    p.add_argument(
        "directory",
        nargs="?",
        type=Path,
        default=Path.cwd(),
        help="Root directory (default: cwd)",
    )
    p.add_argument("--ignore", nargs="*", default=[], help="Extra ignore patterns")
    p.add_argument("--max-depth", type=int, help="Limit recursion depth")
    p.add_argument("--show-sizes", action="store_true", help="Show file sizes")
    p.add_argument(
        "--files-first",
        action="store_true",
        help="List files before folders (default: folders first)",
    )
    p.add_argument("--reverse", action="store_true", help="Reverse sort order")


def build_cli(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(
        prog="mydir",
        description="Directory tree visualiser & bundler",
        epilog=(
            "Use one of the subcommands and add --help for details.\n"
            "Examples:\n"
            "  mydir tree \n"
            "  mydir bundle \n"
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    sub = parser.add_subparsers(dest="cmd", required=True)

    # ---- tree ----
    tree_p = sub.add_parser("tree", help="Print or save a directory tree")
    _add_common_tree_args(tree_p)
    tree_p.add_argument("-o", "--output", type=Path, help="Save to file (.txt or .md)")

    # ---- bundle ----
    bundle_p = sub.add_parser("bundle", help="Create code bundle (PDF/DOCX)")
    _add_common_tree_args(bundle_p)
    bundle_p.add_argument("--pdf", type=Path, help="Target PDF file")
    bundle_p.add_argument("--docx", type=Path, help="Target DOCX file")
    bundle_p.add_argument("--pdf-page-size", default=PDF_DEFAULT_PAGE_SIZE)
    bundle_p.add_argument("--pdf-dpi", type=int, default=PDF_DEFAULT_DPI)
    bundle_p.add_argument("--pdf-scale", type=float, default=PDF_DEFAULT_SCALE)

    args = parser.parse_args(argv)

    tree = (
        mydir(root=args.directory)
        .ignore(*args.ignore)
        .max_depth(args.max_depth)
        .show_sizes_on(args.show_sizes)
    )
    if args.files_first:
        tree.files_first()
    if args.reverse:
        tree.reverse = True

    if args.cmd == "tree":
        if args.output:
            tree.save(args.output)
            print(f"[mydir] Written to {args.output}")
        else:
            tree.print()
    elif args.cmd == "bundle":
        if not args.pdf and not args.docx:
            parser.error("--pdf and/or --docx required for 'bundle'")
        if args.pdf:
            tree.export_code_pdf(
                args.pdf,
                page_size=args.pdf_page_size,
                dpi=args.pdf_dpi,
                scale=args.pdf_scale,
            )
            print(f"[mydir] PDF bundle saved to {args.pdf}")
        if args.docx:
            tree.export_code_docx(args.docx)
            print(f"[mydir] DOCX bundle saved to {args.docx}")


# ---------------------------------------------------------------------------#
# `python -m mydir` entry
# ---------------------------------------------------------------------------#
def _main() -> None:  # pragma: no cover
    build_cli()


if __name__ == "__main__":  # pragma: no cover
    _main()
