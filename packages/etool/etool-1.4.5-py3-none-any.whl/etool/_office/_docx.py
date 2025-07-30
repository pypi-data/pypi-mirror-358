import docx
import os
import re
import cv2
import numpy as np

class ManagerDocx:
    @staticmethod
    def replace_words(path: str, old: str, new: str) -> None:
        """
        Replace keywords in a Word document.


        :param path: file path
        :param old: keyword to replace
        :param new: new keyword after replacement
        :return: None
        """
        if path.endswith(".docx"):

            # does not support reading doc format files
            doc = docx.Document(path)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:

                    if run.text:
                        run.text = run.text.replace(old, new)
                doc.save(path)
        else:
            raise ValueError("Only docx file format is supported!")
        return path


    @staticmethod
    def change_forward(word_path: str, save_path: str) -> None:
        """
        Change the page direction of a Word document.


        :param word_path: Word file path
        :param save_path: result file path
        :return: None
        """
        doc = docx.Document(word_path)
        for section in doc.sections:

            # alternate width and height
            section.page_width, section.page_height = section.page_height, section.page_width
        # save as a new file
        doc.save(save_path) 
        return save_path


    @staticmethod
    def get_pictures(word_path: str, result_path: str) -> str:
        """
        Extract images from a Word document and save them.


        :param word_path: Word file path
        :param result_path: image save path
        :return: image save path
        """
        # create save path
        if not os.path.exists(result_path):

            os.makedirs(result_path)
        # read file
        doc = docx.Document(word_path)


        # get images
        dict_rel = doc.part._rels
        for rel in dict_rel:
            rel = dict_rel[rel]

            if "image" in rel.target_ref:            
                img_name = re.findall("/(.*)", rel.target_ref)[0]
                word_name = os.path.splitext(word_path)[0]
                if os.sep in word_name:
                    new_name = word_name.split('\\')[-1]
                else:
                    new_name = word_name.split('/')[-1]
                # get image size with cv2
                imgdata = np.frombuffer(rel.target_part.blob,np.uint8)
                img_cv = cv2.imdecode(imgdata,cv2.IMREAD_ANYCOLOR)
                img_name = '{}-{}-{}-{}'.format(new_name,img_cv.shape[0],img_cv.shape[1],img_name)

                # write directly to binary for better compatibility than using CV2 to save images
                with open(f'{result_path}/{img_name}','wb') as f:
                    f.write(rel.target_part.blob)
            else:

                pass
        return result_path
    
