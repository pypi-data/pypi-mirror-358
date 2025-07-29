import os
from docx import Document
from docx.shared import Pt
import markdown
from bs4 import BeautifulSoup
import openpyxl
import sys
from pathlib import Path

class ManagerMd:
    """
    管理Markdown文件与Word文档之间的转换，以及其他Markdown转换功能
    """
    
    @staticmethod
    def convert_md_to_docx(md_path, docx_path):
        """
        将Markdown文件转换为Word文档
        
        Args:
            md_path: Markdown文件路径
            docx_path: 要保存的Word文档路径
            
        Returns:
            str: 成功信息
        """
        # 确保输入文件存在
        if not os.path.exists(md_path):
            raise FileNotFoundError(f"找不到Markdown文件: {md_path}")
        
        # 读取Markdown文件内容
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 将Markdown转换为HTML
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        
        # 使用BeautifulSoup解析HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 创建Word文档
        doc = Document()
        
        # 处理HTML内容并添加到Word文档
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'a', 'strong', 'em', 'code', 'pre', 'blockquote']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(element.get_text())
                run.bold = True
                run.font.size = Pt(20 - level)  # h1最大，依次减小
                paragraph.style = f'Heading {level}'
            
            elif element.name == 'p':
                paragraph = doc.add_paragraph()
                ManagerMd._process_paragraph_content(paragraph, element)
            
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    paragraph = doc.add_paragraph(style='List Bullet')
                    ManagerMd._process_paragraph_content(paragraph, li)
            
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    paragraph = doc.add_paragraph(style='List Number')
                    ManagerMd._process_paragraph_content(paragraph, li)
                    
            elif element.name == 'blockquote':
                paragraph = doc.add_paragraph()
                paragraph.style = 'Intense Quote'
                ManagerMd._process_paragraph_content(paragraph, element)

        # 确保输出目录存在
        output_dir = os.path.dirname(docx_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 保存Word文档
        doc.save(docx_path)
        return f"已将Markdown文件转换为Word文档并保存至: {docx_path}"

    @staticmethod
    def _process_paragraph_content(paragraph, element):
        """
        处理段落内容，包括加粗、斜体、链接等格式
        
        Args:
            paragraph: Word文档段落对象
            element: HTML元素
        """
        # 如果元素只包含文本，没有其他格式
        if len(element.contents) == 1 and isinstance(element.contents[0], str):
            paragraph.add_run(element.get_text())
            return
        
        # 递归处理混合格式
        for content in element.contents:
            if isinstance(content, str):
                paragraph.add_run(content)
            else:
                if content.name == 'strong' or content.name == 'b':
                    run = paragraph.add_run(content.get_text())
                    run.bold = True
                elif content.name == 'em' or content.name == 'i':
                    run = paragraph.add_run(content.get_text())
                    run.italic = True
                elif content.name == 'a':
                    run = paragraph.add_run(content.get_text())
                    run.underline = True
                    # 添加超链接
                    run.hyperlink = content.get('href', '')
                elif content.name == 'code':
                    run = paragraph.add_run(content.get_text())
                    run.font.name = 'Courier New'
                else:
                    # 递归处理嵌套元素
                    ManagerMd._process_paragraph_content(paragraph, content)

    @staticmethod
    def convert_md_to_html(md_path, html_path):
        """
        将Markdown文件转换为HTML网页
        
        Args:
            md_path: Markdown文件路径
            html_path: 要保存的HTML文件路径
            
        Returns:
            str: 成功信息
        """
        # 确保输入文件存在
        if not os.path.exists(md_path):
            raise FileNotFoundError(f"找不到Markdown文件: {md_path}")
        
        # 读取Markdown文件内容
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 将Markdown转换为HTML
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        
        # 添加CSS样式
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Markdown转HTML</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                    color: #333;
                }}
                h1, h2, h3, h4, h5, h6 {{
                    color: #444;
                    margin-top: 24px;
                    margin-bottom: 16px;
                    font-weight: 600;
                }}
                h1 {{ font-size: 2em; border-bottom: 1px solid #eaecef; padding-bottom: .3em; }}
                h2 {{ font-size: 1.5em; border-bottom: 1px solid #eaecef; padding-bottom: .3em; }}
                pre {{
                    background-color: #f6f8fa;
                    border-radius: 3px;
                    padding: 16px;
                    overflow: auto;
                }}
                code {{
                    font-family: 'Courier New', Courier, monospace;
                    background-color: #f6f8fa;
                    padding: 0.2em 0.4em;
                    border-radius: 3px;
                }}
                pre code {{
                    background-color: transparent;
                    padding: 0;
                }}
                blockquote {{
                    border-left: 4px solid #dfe2e5;
                    color: #6a737d;
                    padding: 0 1em;
                    margin: 0;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin-bottom: 16px;
                }}
                table th, table td {{
                    border: 1px solid #dfe2e5;
                    padding: 6px 13px;
                }}
                table th {{
                    background-color: #f6f8fa;
                }}
                a {{
                    color: #0366d6;
                    text-decoration: none;
                }}
                a:hover {{
                    text-decoration: underline;
                }}
                img {{
                    max-width: 100%;
                }}
                ul, ol {{
                    padding-left: 2em;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # 确保输出目录存在
        output_dir = os.path.dirname(html_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 保存HTML文件
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(styled_html)
        
        return f"已将Markdown文件转换为HTML网页并保存至: {html_path}"

    @staticmethod
    def extract_tables_to_excel(md_path, excel_path):
        """
        从Markdown文件中提取表格并转换为Excel文件
        
        Args:
            md_path: Markdown文件路径
            excel_path: 要保存的Excel文件路径
            
        Returns:
            str: 成功信息
        """
        # 确保输入文件存在
        if not os.path.exists(md_path):
            raise FileNotFoundError(f"找不到Markdown文件: {md_path}")
        
        # 读取Markdown文件内容
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 将Markdown转换为HTML
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        
        # 使用BeautifulSoup解析HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 查找所有表格
        tables = soup.find_all('table')
        
        if not tables:
            return f"未在Markdown文件中找到表格: {md_path}"
        
        # 确保输出目录存在
        output_dir = os.path.dirname(excel_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 创建Excel工作簿
        workbook = openpyxl.Workbook()
        
        # 删除默认创建的sheet
        default_sheet = workbook.active
        workbook.remove(default_sheet)
        
        # 处理每个表格
        for i, table in enumerate(tables):
            # 创建新的工作表
            sheet_name = f"Table{i+1}"
            sheet = workbook.create_sheet(sheet_name)
            
            # 提取表头
            headers = []
            header_row = table.find('thead')
            if header_row:
                th_elements = header_row.find_all('th')
                headers = [th.get_text(strip=True) for th in th_elements]
                
                # 写入表头
                for col_idx, header in enumerate(headers, 1):
                    sheet.cell(row=1, column=col_idx, value=header)
            
            # 提取表格内容
            rows = table.find_all('tr')
            start_row = 1 if not header_row else 2
            
            for row_idx, row in enumerate(rows, start_row):
                # 跳过表头行
                if header_row and row_idx == 1:
                    continue
                
                # 提取单元格
                cells = row.find_all(['td', 'th'])
                for col_idx, cell in enumerate(cells, 1):
                    value = cell.get_text(strip=True)
                    sheet.cell(row=row_idx, column=col_idx, value=value)
        
        # 保存Excel文件
        workbook.save(excel_path)
        
        return f"已将Markdown文件中的表格提取并转换为Excel文件: {excel_path}"

