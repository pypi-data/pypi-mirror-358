import os
import sys
import logging
import time
import json
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Union, List, Optional, Dict, Callable
from concurrent.futures import ThreadPoolExecutor
from functools import wraps, lru_cache
import argparse
from datetime import datetime

# 第三方库导入
try:
    from pdfminer.high_level import extract_text
    from pdf2image import convert_from_path
    from docx import Document
    import pandas as pd
    from tabula import read_pdf
    from PIL import Image
    from PyPDF2 import PdfReader
except ImportError as e:
    print(f"缺少依赖库: {e}")
    sys.exit(1)

# 日志配置
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('pdf_converter.log')
    ]
)

# 自定义异常
class ConversionError(Exception):
    """基础转换异常"""
    pass

class PDFPermissionError(ConversionError):
    """PDF权限错误"""
    pass

class PDFCorruptedError(ConversionError):
    """PDF文件损坏"""
    pass

class UnsupportedFormatError(ConversionError):
    """不支持的格式"""
    pass

class PDFConverter(ABC):
    """
    PDF文件转换器抽象基类
    
    提供PDF文件转换的基础功能，包括:
    - 文件路径验证
    - 输出路径处理
    - 基本错误处理
    
    子类应实现:
    - convert() 方法: 执行实际转换逻辑
    - supported_formats() 类方法: 返回支持的格式列表
    """
    
    def __init__(self, pdf_path: Union[str, Path]):
        """
        初始化PDF转换器
        
        参数:
            pdf_path: PDF文件路径
        """
        self.pdf_path = Path(pdf_path)
        self.logger = logging.getLogger(self.__class__.__name__)
        self._progress_callback = None
        
        # 验证文件
        if not self.pdf_path.exists():
            raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
        if self.pdf_path.suffix.lower() != '.pdf':
            raise ValueError("输入文件必须是PDF格式")
        
        self.logger.info(f"初始化转换器，处理文件: {pdf_path}")

    @abstractmethod
    def convert(self, output_path: Union[str, Path], **kwargs) -> Union[str, List[str]]:
        """
        将PDF转换为目标格式
        
        参数:
            output_path: 输出文件路径
            **kwargs: 转换选项
            
        返回:
            转换后的文件路径或路径列表
        """
        pass

    @classmethod
    @abstractmethod
    def supported_formats(cls) -> List[str]:
        """返回支持的格式列表"""
        return []

    def set_progress_callback(self, callback: Callable[[int, int], None]):
        """设置进度回调函数"""
        self._progress_callback = callback
        
    def _update_progress(self, current: int, total: int):
        """更新进度"""
        if self._progress_callback:
            self._progress_callback(current, total)

    def get_metadata(self) -> Dict[str, str]:
        """获取PDF元数据"""
        try:
            with open(self.pdf_path, 'rb') as f:
                reader = PdfReader(f)
                return {
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'pages': len(reader.pages),
                    'created': reader.metadata.get('/CreationDate', ''),
                    'modified': reader.metadata.get('/ModDate', '')
                }
        except Exception as e:
            self.logger.warning(f"获取元数据失败: {str(e)}")
            return {}

    def _prepare_output_path(self, output_path: Union[str, Path], 
                           default_extension: str) -> Path:
        """
        准备输出路径
        
        参数:
            output_path: 输出路径
            default_extension: 默认文件扩展名
            
        返回:
            处理后的Path对象
        """
        output_path = Path(output_path)
        
        # 如果是目录，自动生成文件名
        if output_path.is_dir():
            output_path = output_path / f"{self.pdf_path.stem}.{default_extension}"
        # 如果没有扩展名，添加默认扩展名
        elif not output_path.suffix:
            output_path = output_path.with_suffix(f".{default_extension}")
            
        # 创建父目录（如果不存在）
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        return output_path

# 重试装饰器
def retry(max_attempts=3, delay=1, exceptions=(Exception,)):
    def decorator(f):
        @wraps(f)
        def wrapper(*args, **kwargs):
            last_error = None
            for attempt in range(1, max_attempts+1):
                try:
                    return f(*args, **kwargs)
                except exceptions as e:
                    last_error = e
                    if attempt < max_attempts:
                        time.sleep(delay)
            raise last_error
        return wrapper
    return decorator

# 具体转换器实现
class PDFToDocxConverter(PDFConverter):
    @retry(max_attempts=3, delay=0.5)
    def convert(self, output_path: Union[str, Path], **kwargs) -> str:
        """
        将PDF转换为Word文档(.docx)
        
        参数:
            output_path: 输出文件路径
            **kwargs:
                - start_page: 开始页(从1开始)
                - end_page: 结束页
                - preserve_formatting: 是否保留格式(True/False)
                
        返回:
            转换后的文件路径
        """
        try:
            output_path = self._prepare_output_path(output_path, 'docx')
            
            start_page = kwargs.get('start_page', 1)
            end_page = kwargs.get('end_page', None)
            preserve = kwargs.get('preserve_formatting', False)
            
            # 尝试使用pdf2docx库（如果安装）
            try:
                from pdf2docx import Converter
                cv = Converter(str(self.pdf_path))
                cv.convert(str(output_path), start=start_page, end=end_page)
                cv.close()
                return str(output_path)
            except ImportError:
                self.logger.warning("pdf2docx未安装，使用基本文本转换")
                # 回退到基本实现
                doc = Document()
                text = extract_text(
                    str(self.pdf_path),
                    page_numbers=range(start_page-1, end_page) if end_page else None
                )
                
                for paragraph in text.split('\n\n'):
                    if paragraph.strip():
                        para = doc.add_paragraph()
                        if preserve:
                            runs = paragraph.split('\n')
                            for run in runs:
                                if run.strip():
                                    para.add_run(run.strip() + ' ')
                        else:
                            para.add_run(paragraph.strip())
                
                doc.save(output_path)
                return str(output_path)
        except Exception as e:
            raise ConversionError(f"转换为DOCX失败: {str(e)}")

    @classmethod
    def supported_formats(cls) -> List[str]:
        return ['docx']

class PDFToImageConverter(PDFConverter):
    @retry(max_attempts=3, delay=0.5)
    def convert(self, output_path: Union[str, Path], **kwargs) -> Union[str, List[str]]:
        """
        将PDF转换为图像
        
        参数:
            output_path: 输出文件路径或目录
            **kwargs:
                - dpi: 图像DPI(默认200)
                - fmt: 图像格式(png/jpg/tiff)
                - merge: 是否合并所有页为一张长图(True/False)
                - quality: 图像质量(1-100)
                
        返回:
            单个文件路径或多个文件路径列表
        """
        try:
            dpi = kwargs.get('dpi', 200)
            fmt = kwargs.get('fmt', 'png').lower()
            merge = kwargs.get('merge', False)
            quality = kwargs.get('quality', 90)
            
            if fmt not in ['png', 'jpg', 'jpeg', 'tiff']:
                raise ValueError(f"不支持的图像格式: {fmt}")
                
            images = convert_from_path(
                str(self.pdf_path),
                dpi=dpi,
                fmt=fmt if fmt != 'jpg' else 'jpeg'
            )
            
            self._update_progress(0, len(images))
            
            if merge:
                # 合并所有页为一张长图
                output_path = self._prepare_output_path(output_path, fmt)
                total_height = sum(img.height for img in images)
                max_width = max(img.width for img in images)
                
                merged_image = Image.new('RGB', (max_width, total_height))
                y_offset = 0
                for i, img in enumerate(images):
                    merged_image.paste(img, (0, y_offset))
                    y_offset += img.height
                    self._update_progress(i+1, len(images))
                
                merged_image.save(output_path, quality=quality)
                return str(output_path)
            else:
                # 每页保存为单独图像
                output_path = Path(output_path)
                if len(images) == 1:
                    output_path = self._prepare_output_path(output_path, fmt)
                    images[0].save(output_path, quality=quality)
                    self._update_progress(1, 1)
                    return str(output_path)
                else:
                    output_path.mkdir(parents=True, exist_ok=True)
                    output_files = []
                    
                    # 多页并行处理
                    if len(images) > 5:
                        with ThreadPoolExecutor() as executor:
                            results = list(executor.map(
                                lambda x: self._save_image(x[1], output_path/f"page_{x[0]+1}.{fmt}", quality),
                                enumerate(images)
                            ))
                            output_files.extend(results)
                    else:
                        for i, image in enumerate(images):
                            page_path = output_path / f"page_{i+1}.{fmt}"
                            image.save(page_path, quality=quality)
                            output_files.append(str(page_path))
                            self._update_progress(i+1, len(images))
                    
                    return output_files
        except Exception as e:
            raise ConversionError(f"转换为图像失败: {str(e)}")

    def _save_image(self, image: Image.Image, path: Path, quality: int) -> str:
        """保存单个图像"""
        image.save(path, quality=quality)
        return str(path)

    @classmethod
    def supported_formats(cls) -> List[str]:
        return ['png', 'jpg', 'jpeg', 'tiff']

class PDFToTextConverter(PDFConverter):
    @retry(max_attempts=3, delay=0.5)
    def convert(self, output_path: Union[str, Path], **kwargs) -> str:
        """
        将PDF转换为纯文本
        
        参数:
            output_path: 输出文件路径
            **kwargs:
                - start_page: 开始页(从1开始)
                - end_page: 结束页
                - encoding: 文本编码(默认utf-8)
                
        返回:
            转换后的文件路径
        """
        try:
            output_path = self._prepare_output_path(output_path, 'txt')
            
            start_page = kwargs.get('start_page', 1)
            end_page = kwargs.get('end_page', None)
            encoding = kwargs.get('encoding', 'utf-8')
            
            text = extract_text(
                str(self.pdf_path),
                page_numbers=range(start_page-1, end_page) if end_page else None
            )
            
            with open(output_path, 'w', encoding=encoding) as f:
                f.write(text)
            
            return str(output_path)
        except Exception as e:
            raise ConversionError(f"转换为文本失败: {str(e)}")

    @classmethod
    def supported_formats(cls) -> List[str]:
        return ['txt']

class PDFToCSVConverter(PDFConverter):
    @retry(max_attempts=3, delay=0.5)
    def convert(self, output_path: Union[str, Path], **kwargs) -> Union[str, List[str]]:
        """
        提取PDF中的表格为CSV
        
        参数:
            output_path: 输出文件路径或目录
            **kwargs:
                - pages: 要提取的页码('all'或数字或列表)
                - multiple_tables: 如何处理多个表格(separate/merge)
                - encoding: CSV文件编码(默认utf-8)
                
        返回:
            单个CSV文件路径或多个CSV文件路径列表
        """
        try:
            pages = kwargs.get('pages', 'all')
            multiple_tables = kwargs.get('multiple_tables', 'separate')
            encoding = kwargs.get('encoding', 'utf-8')
            
            dfs = read_pdf(str(self.pdf_path), pages=pages, multiple_tables=True)
            
            if not dfs:
                raise ConversionError("未找到表格数据")
                
            self._update_progress(0, len(dfs))
                
            if multiple_tables == 'merge':
                # 合并所有表格
                output_path = self._prepare_output_path(output_path, 'csv')
                merged_df = pd.concat(dfs, ignore_index=True)
                merged_df.to_csv(output_path, index=False, encoding=encoding)
                self._update_progress(1, 1)
                return str(output_path)
            else:
                # 每个表格保存为单独CSV
                output_path = Path(output_path)
                if len(dfs) == 1:
                    output_path = self._prepare_output_path(output_path, 'csv')
                    dfs[0].to_csv(output_path, index=False, encoding=encoding)
                    self._update_progress(1, 1)
                    return str(output_path)
                else:
                    output_path.mkdir(parents=True, exist_ok=True)
                    output_files = []
                    for i, df in enumerate(dfs):
                        table_path = output_path / f"table_{i+1}.csv"
                        df.to_csv(table_path, index=False, encoding=encoding)
                        output_files.append(str(table_path))
                        self._update_progress(i+1, len(dfs))
                    return output_files
        except Exception as e:
            raise ConversionError(f"提取表格失败: {str(e)}")

    @classmethod
    def supported_formats(cls) -> List[str]:
        return ['csv']

class PDFToHTMLConverter(PDFConverter):
    @retry(max_attempts=3, delay=0.5)
    def convert(self, output_path: Union[str, Path], **kwargs) -> str:
        """
        将PDF转换为HTML
        
        参数:
            output_path: 输出文件路径
            **kwargs:
                - css: 自定义CSS样式
                - images: 是否嵌入图像(True/False)
                
        返回:
            转换后的HTML文件路径
        """
        try:
            output_path = self._prepare_output_path(output_path, 'html')
            
            # 使用pdfminer提取文本
            text = extract_text(str(self.pdf_path))
            
            # 简单的HTML转换
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{self.pdf_path.stem}</title>
    <style>
        {kwargs.get('css', 'body {{ font-family: Arial; margin: 20px; }}')}
    </style>
</head>
<body>
    <h1>{self.pdf_path.stem}</h1>
    <div id="content">
        {text.replace('\n', '<br>')}
    </div>
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return str(output_path)
        except Exception as e:
            raise ConversionError(f"转换为HTML失败: {str(e)}")

    @classmethod
    def supported_formats(cls) -> List[str]:
        return ['html']

# 工厂类
class PDFConverterFactory:
    _format_map = {
        'docx': (PDFToDocxConverter, 'Microsoft Word文档'),
        'txt': (PDFToTextConverter, '纯文本文件'),
        'png': (PDFToImageConverter, 'PNG图像'),
        'jpg': (PDFToImageConverter, 'JPEG图像'),
        'jpeg': (PDFToImageConverter, 'JPEG图像'),
        'tiff': (PDFToImageConverter, 'TIFF图像'),
        'csv': (PDFToCSVConverter, 'CSV表格数据'),
        'html': (PDFToHTMLConverter, 'HTML网页'),
    }

    @staticmethod
    @lru_cache(maxsize=32)
    def get_converter(target_format: str, pdf_path: str) -> PDFConverter:
        """
        获取指定格式的转换器
        
        参数:
            target_format: 目标格式
            pdf_path: PDF文件路径
            
        返回:
            PDFConverter实例
            
        抛出:
            UnsupportedFormatError: 当格式不支持时
        """
        target_format = target_format.lower()
        if target_format not in PDFConverterFactory._format_map:
            raise UnsupportedFormatError(f"不支持的格式: {target_format}")
        
        return PDFConverterFactory._format_map[target_format][0](pdf_path)

    @staticmethod
    def get_supported_formats() -> Dict[str, str]:
        """获取所有支持的格式及其描述"""
        return {fmt: desc for fmt, (_, desc) in PDFConverterFactory._format_map.items()}

# 命令行接口
def parse_args():
    parser = argparse.ArgumentParser(
        description='PDF转换工具 - 支持多种格式转换',
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument('input', help='输入PDF文件路径')
    parser.add_argument('output', help='输出文件路径或目录')
    parser.add_argument('-f', '--format', required=True, 
                       choices=PDFConverterFactory.get_supported_formats().keys(),
                       help='目标格式')
    parser.add_argument('--dpi', type=int, default=200, 
                       help='图像DPI(仅图像转换)')
    parser.add_argument('--start-page', type=int, default=1,
                       help='起始页码(从1开始)')
    parser.add_argument('--end-page', type=int, 
                       help='结束页码(默认为最后一页)')
    parser.add_argument('--preserve-formatting', action='store_true',
                       help='保留格式(仅DOCX转换)')
    parser.add_argument('--merge', action='store_true',
                       help='合并多页为单个文件(图像/表格)')
    parser.add_argument('--quality', type=int, default=90,
                       help='输出质量(1-100, 仅图像)')
    return parser.parse_args()

def main():
    args = parse_args()
    
    try:
        # 获取转换器
        converter = PDFConverterFactory.get_converter(args.format, args.input)
        
        # 设置进度回调
        def progress_callback(current, total):
            print(f"\r进度: {current}/{total} ({current/total:.1%})", end='', flush=True)
        converter.set_progress_callback(progress_callback)
        
        # 执行转换
        print(f"开始转换: {args.input} -> {args.output} ({args.format})")
        result = converter.convert(
            args.output,
            dpi=args.dpi,
            start_page=args.start_page,
            end_page=args.end_page,
            preserve_formatting=args.preserve_formatting,
            merge=args.merge,
            quality=args.quality
        )
        
        print("\n转换成功!")
        if isinstance(result, list):
            print(f"生成文件 {len(result)} 个:")
            for file in result:
                print(f"  - {file}")
        else:
            print(f"输出文件: {result}")
            
    except Exception as e:
        print(f"\n转换失败: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()