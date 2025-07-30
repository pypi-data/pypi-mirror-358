import os
import sys
import logging
import time
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Union, List, Optional, Dict, Callable
from concurrent.futures import ThreadPoolExecutor
from functools import wraps, lru_cache
import argparse

# 第三方库导入
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from PIL import Image
    import pandas as pd
except ImportError as e:
    print(f"缺少依赖库: {e}")
    sys.exit(1)

# 日志配置
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('docx_processor.log')
    ]
)

# 自定义异常
class DocxProcessingError(Exception):
    """基础处理异常"""
    pass

class DocxPermissionError(DocxProcessingError):
    """DOCX权限错误"""
    pass

class DocxCorruptedError(DocxProcessingError):
    """DOCX文件损坏"""
    pass

class UnsupportedFormatError(DocxProcessingError):
    """不支持的格式"""
    pass

class DocxProcessor(ABC):
    """
    DOCX文档处理器抽象基类
    
    提供DOCX文档处理的基础功能，包括:
    - 文件路径验证
    - 输出路径处理
    - 基本错误处理
    
    子类应实现:
    - process() 方法: 执行实际处理逻辑
    - supported_operations() 类方法: 返回支持的操作列表
    """
    
    def __init__(self, docx_path: Union[str, Path]):
        """
        初始化DOCX处理器
        
        参数:
            docx_path: DOCX文件路径
        """
        self.docx_path = Path(docx_path)
        self.logger = logging.getLogger(self.__class__.__name__)
        self._progress_callback = None
        
        # 验证文件
        if not self.docx_path.exists():
            raise FileNotFoundError(f"DOCX文件不存在: {docx_path}")
        if self.docx_path.suffix.lower() != '.docx':
            raise ValueError("输入文件必须是DOCX格式")
        
        self.logger.info(f"初始化处理器，处理文件: {docx_path}")

    @abstractmethod
    def process(self, output_path: Union[str, Path], **kwargs) -> Union[str, List[str]]:
        """
        处理DOCX文档
        
        参数:
            output_path: 输出文件路径
            **kwargs: 处理选项
            
        返回:
            处理后的文件路径或路径列表
        """
        pass

    @classmethod
    @abstractmethod
    def supported_operations(cls) -> List[str]:
        """返回支持的操作列表"""
        return []

    def set_progress_callback(self, callback: Callable[[int, int], None]):
        """设置进度回调函数"""
        self._progress_callback = callback
        
    def _update_progress(self, current: int, total: int):
        """更新进度"""
        if self._progress_callback:
            self._progress_callback(current, total)

    def get_metadata(self) -> Dict[str, str]:
        """获取DOCX元数据"""
        try:
            doc = Document(self.docx_path)
            return {
                'title': doc.core_properties.title,
                'author': doc.core_properties.author,
                'created': str(doc.core_properties.created),
                'modified': str(doc.core_properties.modified),
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
        except Exception as e:
            self.logger.warning(f"获取元数据失败: {str(e)}")
            return {}

    def _prepare_output_path(self, output_path: Union[str, Path], 
                           default_extension: str) -> Path:
        """
        准备输出路径
        
        参数:
            output_path: 输出路径
            default_extension: 默认文件扩展名
            
        返回:
            处理后的Path对象
        """
        output_path = Path(output_path)
        
        # 如果是目录，自动生成文件名
        if output_path.is_dir():
            output_path = output_path / f"{self.docx_path.stem}.{default_extension}"
        # 如果没有扩展名，添加默认扩展名
        elif not output_path.suffix:
            output_path = output_path.with_suffix(f".{default_extension}")
            
        # 创建父目录（如果不存在）
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        return output_path

# 重试装饰器
def retry(max_attempts=3, delay=1, exceptions=(Exception,)):
    def decorator(f):
        @wraps(f)
        def wrapper(*args, **kwargs):
            last_error = None
            for attempt in range(1, max_attempts+1):
                try:
                    return f(*args, **kwargs)
                except exceptions as e:
                    last_error = e
                    if attempt < max_attempts:
                        time.sleep(delay)
            raise last_error
        return wrapper
    return decorator

# 具体处理器实现
class DocxToPDFConverter(DocxProcessor):
    @retry(max_attempts=3, delay=0.5)
    def process(self, output_path: Union[str, Path], **kwargs) -> str:
        """
        将DOCX转换为PDF
        
        参数:
            output_path: 输出文件路径
            **kwargs:
                - pages: 要转换的页码范围(如'1-3,5')
                - quality: 输出质量(high/medium/low)
                
        返回:
            转换后的文件路径
        """
        try:
            output_path = self._prepare_output_path(output_path, 'pdf')
            
            # 尝试使用docx2pdf库
            try:
                from docx2pdf import convert
                convert(str(self.docx_path), str(output_path))
                return str(output_path)
            except ImportError:
                self.logger.warning("docx2pdf未安装，尝试使用unoconv")
                # 回退到unoconv
                try:
                    import subprocess
                    subprocess.run(['unoconv', '-f', 'pdf', '-o', str(output_path), str(self.docx_path)], check=True)
                    return str(output_path)
                except Exception as e:
                    raise DocxProcessingError(f"转换为PDF失败，请安装docx2pdf或unoconv: {str(e)}")
        except Exception as e:
            raise DocxProcessingError(f"转换为PDF失败: {str(e)}")

    @classmethod
    def supported_operations(cls) -> List[str]:
        return ['pdf']

class DocxToTextConverter(DocxProcessor):
    @retry(max_attempts=3, delay=0.5)
    def process(self, output_path: Union[str, Path], **kwargs) -> str:
        """
        将DOCX转换为纯文本
        
        参数:
            output_path: 输出文件路径
            **kwargs:
                - include_tables: 是否包含表格内容(True/False)
                - encoding: 文本编码(默认utf-8)
                
        返回:
            转换后的文件路径
        """
        try:
            output_path = self._prepare_output_path(output_path, 'txt')
            include_tables = kwargs.get('include_tables', False)
            encoding = kwargs.get('encoding', 'utf-8')
            
            doc = Document(self.docx_path)
            text = []
            
            for para in doc.paragraphs:
                text.append(para.text)
            
            if include_tables:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text.append(cell.text)
            
            with open(output_path, 'w', encoding=encoding) as f:
                f.write('\n'.join(text))
            
            return str(output_path)
        except Exception as e:
            raise DocxProcessingError(f"转换为文本失败: {str(e)}")

    @classmethod
    def supported_operations(cls) -> List[str]:
        return ['txt']

class DocxToHTMLConverter(DocxProcessor):
    @retry(max_attempts=3, delay=0.5)
    def process(self, output_path: Union[str, Path], **kwargs) -> str:
        """
        将DOCX转换为HTML
        
        参数:
            output_path: 输出文件路径
            **kwargs:
                - css: 自定义CSS样式
                - include_images: 是否包含图像(True/False)
                
        返回:
            转换后的HTML文件路径
        """
        try:
            output_path = self._prepare_output_path(output_path, 'html')
            
            # 尝试使用pandoc
            try:
                import subprocess
                subprocess.run([
                    'pandoc', '-s', str(self.docx_path), 
                    '-o', str(output_path), 
                    '--css', kwargs.get('css', '')
                ], check=True)
                return str(output_path)
            except Exception:
                self.logger.warning("pandoc不可用，使用基本转换")
                # 基本实现
                doc = Document(self.docx_path)
                html_content = [
                    '<!DOCTYPE html>',
                    '<html>',
                    '<head>',
                    '<meta charset="UTF-8">',
                    f'<title>{self.docx_path.stem}</title>',
                    f'<style>{kwargs.get("css", "body { font-family: Arial; }")}</style>',
                    '</head>',
                    '<body>'
                ]
                
                for para in doc.paragraphs:
                    html_content.append(f'<p>{para.text}</p>')
                
                html_content.extend(['</body>', '</html>'])
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(html_content))
                
                return str(output_path)
        except Exception as e:
            raise DocxProcessingError(f"转换为HTML失败: {str(e)}")

    @classmethod
    def supported_operations(cls) -> List[str]:
        return ['html']

class DocxToCSVConverter(DocxProcessor):
    @retry(max_attempts=3, delay=0.5)
    def process(self, output_path: Union[str, Path], **kwargs) -> Union[str, List[str]]:
        """
        提取DOCX中的表格为CSV
        
        参数:
            output_path: 输出文件路径或目录
            **kwargs:
                - table_indexes: 要提取的表格索引列表(如[0,2])
                - encoding: CSV文件编码(默认utf-8)
                
        返回:
            单个CSV文件路径或多个CSV文件路径列表
        """
        try:
            table_indexes = kwargs.get('table_indexes', None)
            encoding = kwargs.get('encoding', 'utf-8')
            
            doc = Document(self.docx_path)
            tables = doc.tables
            
            if not tables:
                raise DocxProcessingError("未找到表格数据")
            
            if table_indexes is None:
                table_indexes = range(len(tables))
            
            output_path = Path(output_path)
            if len(table_indexes) == 1:
                output_path = self._prepare_output_path(output_path, 'csv')
                df = self._table_to_dataframe(tables[table_indexes[0]])
                df.to_csv(output_path, index=False, encoding=encoding)
                return str(output_path)
            else:
                output_path.mkdir(parents=True, exist_ok=True)
                output_files = []
                for i in table_indexes:
                    if i < len(tables):
                        table_path = output_path / f"table_{i}.csv"
                        df = self._table_to_dataframe(tables[i])
                        df.to_csv(table_path, index=False, encoding=encoding)
                        output_files.append(str(table_path))
                return output_files
        except Exception as e:
            raise DocxProcessingError(f"提取表格失败: {str(e)}")
    
    def _table_to_dataframe(self, table) -> pd.DataFrame:
        """将Word表格转换为DataFrame"""
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            data.append(row_data)
        return pd.DataFrame(data)

    @classmethod
    def supported_operations(cls) -> List[str]:
        return ['csv']

class DocxEditor(DocxProcessor):
    """DOCX文档编辑器"""
    @retry(max_attempts=3, delay=0.5)
    def process(self, output_path: Union[str, Path], **kwargs) -> str:
        """
        编辑DOCX文档
        
        参数:
            output_path: 输出文件路径
            **kwargs:
                - replace: 替换字典 {旧文本: 新文本}
                - add_watermark: 水印文本
                - font_size: 字体大小(默认12)
                - font_color: 字体颜色(十六进制，如#FF0000)
                
        返回:
            编辑后的文件路径
        """
        try:
            output_path = self._prepare_output_path(output_path, 'docx')
            
            doc = Document(self.docx_path)
            
            # 文本替换
            if 'replace' in kwargs:
                for old_text, new_text in kwargs['replace'].items():
                    for para in doc.paragraphs:
                        if old_text in para.text:
                            para.text = para.text.replace(old_text, new_text)
            
            # 添加水印
            if 'add_watermark' in kwargs:
                watermark = kwargs['add_watermark']
                font_size = kwargs.get('font_size', 12)
                font_color = kwargs.get('font_color', '#C0C0C0')
                
                for section in doc.sections:
                    header = section.header
                    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    run = paragraph.add_run(watermark)
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor(*self._hex_to_rgb(font_color))
            
            doc.save(output_path)
            return str(output_path)
        except Exception as e:
            raise DocxProcessingError(f"编辑文档失败: {str(e)}")
    
    def _hex_to_rgb(self, hex_color: str) -> tuple:
        """十六进制颜色转RGB"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

    @classmethod
    def supported_operations(cls) -> List[str]:
        return ['edit']

# 工厂类
class DocxProcessorFactory:
    _operation_map = {
        'pdf': (DocxToPDFConverter, 'PDF文档'),
        'txt': (DocxToTextConverter, '纯文本文件'),
        'html': (DocxToHTMLConverter, 'HTML网页'),
        'csv': (DocxToCSVConverter, 'CSV表格数据'),
        'edit': (DocxEditor, '文档编辑')
    }

    @staticmethod
    @lru_cache(maxsize=32)
    def get_processor(operation: str, docx_path: str) -> DocxProcessor:
        """
        获取指定操作的处理器
        
        参数:
            operation: 操作类型
            docx_path: DOCX文件路径
            
        返回:
            DocxProcessor实例
            
        抛出:
            UnsupportedOperationError: 当操作不支持时
        """
        operation = operation.lower()
        if operation not in DocxProcessorFactory._operation_map:
            raise UnsupportedFormatError(f"不支持的操作: {operation}")
        
        return DocxProcessorFactory._operation_map[operation][0](docx_path)

    @staticmethod
    def get_supported_operations() -> Dict[str, str]:
        """获取所有支持的操作及其描述"""
        return {op: desc for op, (_, desc) in DocxProcessorFactory._operation_map.items()}

# 命令行接口
def parse_args():
    parser = argparse.ArgumentParser(
        description='DOCX文档处理工具 - 支持多种格式转换和编辑',
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument('input', help='输入DOCX文件路径')
    parser.add_argument('output', help='输出文件路径或目录')
    parser.add_argument('-o', '--operation', required=True, 
                       choices=DocxProcessorFactory.get_supported_operations().keys(),
                       help='操作类型')
    parser.add_argument('--pages', help='页码范围(仅PDF转换)')
    parser.add_argument('--quality', choices=['high', 'medium', 'low'], 
                       default='medium', help='输出质量(仅PDF转换)')
    parser.add_argument('--include-tables', action='store_true',
                       help='包含表格内容(仅文本转换)')
    parser.add_argument('--table-indexes', help='要提取的表格索引(如0,2,3)')
    parser.add_argument('--replace', nargs=2, action='append',
                       metavar=('OLD', 'NEW'), help='文本替换(编辑操作)')
    parser.add_argument('--watermark', help='水印文本(编辑操作)')
    parser.add_argument('--font-size', type=int, default=12,
                       help='水印字体大小(编辑操作)')
    parser.add_argument('--font-color', default='#C0C0C0',
                       help='水印字体颜色(十六进制，编辑操作)')
    return parser.parse_args()

def main():
    args = parse_args()
    
    try:
        # 获取处理器
        processor = DocxProcessorFactory.get_processor(args.operation, args.input)
        
        # 准备参数
        kwargs = {}
        if args.operation == 'pdf':
            kwargs.update({
                'pages': args.pages,
                'quality': args.quality
            })
        elif args.operation == 'txt':
            kwargs['include_tables'] = args.include_tables
        elif args.operation == 'csv':
            if args.table_indexes:
                kwargs['table_indexes'] = [int(i) for i in args.table_indexes.split(',')]
        elif args.operation == 'edit':
            if args.replace:
                kwargs['replace'] = dict(args.replace)
            if args.watermark:
                kwargs['add_watermark'] = args.watermark
                kwargs['font_size'] = args.font_size
                kwargs['font_color'] = args.font_color
        
        # 设置进度回调
        def progress_callback(current, total):
            print(f"\r进度: {current}/{total} ({current/total:.1%})", end='', flush=True)
        processor.set_progress_callback(progress_callback)
        
        # 执行处理
        print(f"开始处理: {args.input} -> {args.output} ({args.operation})")
        result = processor.process(args.output, **kwargs)
        
        print("\n处理成功!")
        if isinstance(result, list):
            print(f"生成文件 {len(result)} 个:")
            for file in result:
                print(f"  - {file}")
        else:
            print(f"输出文件: {result}")
            
    except Exception as e:
        print(f"\n处理失败: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()