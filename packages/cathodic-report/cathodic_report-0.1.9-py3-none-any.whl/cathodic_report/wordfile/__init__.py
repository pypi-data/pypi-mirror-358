import pathlib
from io import BytesIO

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def create_report(
    project_name,
    data,
    filepath: pathlib.Path = "杂散电流干扰评价报告.docx",
):
    doc = Document()
    # 设置中文字体
    doc.styles["Normal"].font.name = "黑体"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # 修改标题字体
    for i in range(1, 10):  # Assuming you have up to 9 heading levels
        heading_style = doc.styles[f"Heading {i}"]
        heading_style.font.name = "黑体"
        heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # 修改 Title 样式（用于最顶层标题）
    doc.styles["Title"].font.name = "黑体"
    doc.styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # 添加标题
    doc.add_heading("杂散电流干扰评价报告", 0)

    # 1. 干扰情况总览
    doc.add_heading("1. 干扰情况总览", level=1)
    p = doc.add_paragraph(f"测试管道名称：{project_name}")
    p = doc.add_paragraph(f"评价结果：共测试{len(data)}处，结果如下")

    p = doc.add_paragraph("表1 杂散电流评价结果总览")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Ensure center alignment

    # 添加表格
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "序号"
    hdr_cells[1].text = "测试桩编号"
    hdr_cells[2].text = "里程"
    hdr_cells[3].text = "试片面积"
    hdr_cells[4].text = "土壤电阻率"
    hdr_cells[5].text = "直流干扰评价结果"

    for i, item in enumerate(data, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = item["test_pile_number"]
        row_cells[2].text = f"{item['mileage']} km"
        row_cells[3].text = f"{item['test_piece_area']} A/m²"
        row_cells[4].text = f"{item['soil_resistivity']} Ω•m"
        row_cells[5].text = item["dc_interference_result"]

    # 1.1 饼图：直流评价结果占比图
    doc.add_heading("1.1 饼图：直流评价结果占比图", level=2)

    # 这里添加生成饼图的代码
    # ...

    # 1.2 管道沿线干扰分布图
    doc.add_heading("1.2 管道沿线干扰分布图", level=2)

    # 这里添加生成分布图的代码
    # ...

    # 2. 测试桩数据分析结果
    for item in data:
        doc.add_heading(f"2. {item['test_pile_number']}数据分析结果", level=1)
        p = doc.add_paragraph(f"里程：{item['mileage']} km")
        p = doc.add_paragraph(f"试片面积：{item['test_piece_area']} A/m²")
        p = doc.add_paragraph(f"土壤电阻率：{item['soil_resistivity']} Ω•m")
        p = doc.add_paragraph("评判准则：-0.85 V（参考）")

        # 2.1 直流干扰分析结果
        doc.add_heading("2.1 直流干扰分析结果", level=2)

        # 添加直流干扰分析结果表格
        # ...

        # 2.2 交流干扰分析结果
        doc.add_heading("2.2 交流干扰分析结果", level=2)

        # 添加交流干扰分析结果表格
        # ...

        # 2.3 干扰测试结果（图）
        doc.add_heading("2.3 干扰测试结果（图）", level=2)

        # 添加干扰测试结果图
        # ...

    # 保存文档
    doc.save(filepath)


def add_pic(doc: Document, filepath: pathlib.Path):
    img_stream = filepath.read_bytes()
    doc.add_picture(img_stream)


def add_pic_from_plt(doc, save_fig):
    """
    不涉及 disk 读写，直接写入文件
    save_fig: plt.savefig(img_stream, format='png')
    """
    # 将图片保存到内存中
    img_stream = BytesIO()
    save_fig(img_stream)
    # plt.savefig(img_stream, format="png")
    img_stream.seek(0)

    # 将图片添加到文档中
    doc.add_picture(img_stream, width=Inches(6))
    plt.close()
