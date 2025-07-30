"""
Image Builder Module

This module provides the ImageBuilder class for rendering image blocks in Word documents.
It handles image insertion with automatic sizing, DPI calculation, and error handling.
"""

from docx.shared import Inches, RGBColor
from docx.enum.shape import WD_INLINE_SHAPE
from PIL import Image as PILImage
import os
from docxblocks.constants import DEFAULT_EMPTY_VALUE_TEXT, DEFAULT_EMPTY_VALUE_STYLE

class ImageBuilder:
    """
    Builder class for rendering image blocks in Word documents.
    
    This builder handles image insertion with automatic sizing based on DPI,
    support for maximum width/height constraints, and graceful error handling
    for missing or invalid image files.
    
    The builder automatically calculates image dimensions from DPI information
    and supports scaling based on maximum width/height constraints. Invalid
    or missing images are replaced with consistent placeholders.
    """
    
    @staticmethod
    def build(doc, image_path=None, parent=None, index=None, **style_kwargs):
        """
        Build and render an image block in the document.
        
        This method processes the image path, validates file existence,
        calculates appropriate dimensions, and inserts the image into the
        document. Handles errors gracefully with placeholder text.
        
        Args:
            doc: The python-docx Document object
            image_path: Path to the image file (can be None, empty, or invalid)
            parent: The parent XML element where content will be inserted
            index: The insertion index within the parent element
            **style_kwargs: Additional styling options including:
                - max_width: Maximum width constraint (e.g., "4in", "300px")
                - max_height: Maximum height constraint (e.g., "4in", "300px")
                - wrap_text: Text wrapping mode ("inline", "square", "tight", etc.)
                - horizontal_align: Horizontal alignment ("left", "center", "right")
                - vertical_align: Vertical alignment ("top", "middle", "bottom")
                - distance_from_text: Distance from text (e.g., "0.1in", "10px")
        """
        # Validate required parameters
        if parent is None or index is None:
            return
            
        # Handle empty or invalid image path with placeholder
        if not image_path or not image_path.strip() or not os.path.isfile(image_path):
            if hasattr(parent, 'add_paragraph'):
                para = parent.add_paragraph(DEFAULT_EMPTY_VALUE_TEXT)
            else:
                para = doc.add_paragraph(DEFAULT_EMPTY_VALUE_TEXT)
            # Apply placeholder style
            run = para.runs[0]
            run.font.bold = DEFAULT_EMPTY_VALUE_STYLE.get("bold", True)
            if DEFAULT_EMPTY_VALUE_STYLE.get("font_color"):
                run.font.color.rgb = RGBColor.from_string(DEFAULT_EMPTY_VALUE_STYLE["font_color"])
            parent.insert(index, para._element)
            return

        if hasattr(parent, 'add_paragraph'):
            new_para = parent.add_paragraph()
        else:
            new_para = doc.add_paragraph()
        run = new_para.add_run()

        try:
            with PILImage.open(image_path) as img:
                width_px, height_px = img.size
                dpi_x, dpi_y = img.info.get("dpi", (72, 72))
                width_in = width_px / dpi_x
                height_in = height_px / dpi_y

                # Calculate scale factors for width and height constraints
                scales = []
                max_width = _parse_measurement(style_kwargs.get("max_width"))
                max_height = _parse_measurement(style_kwargs.get("max_height"))

                if max_width:
                    scales.append(max_width / width_in)
                if max_height:
                    scales.append(max_height / height_in)
                
                # Use the minimum scale to ensure neither dimension exceeds its maximum
                # This allows upscaling if the image is smaller than the specified constraints
                if scales:
                    scale = min(scales)
                else:
                    scale = 1.0

                # Use file object to ensure embedding works in headers/footers
                with open(image_path, 'rb') as img_file:
                    run.add_picture(img_file, width=Inches(width_in * scale), height=Inches(height_in * scale))

                # Apply text wrapping and positioning properties
                _apply_image_wrapping(run, style_kwargs)

        except Exception as e:
            if hasattr(parent, 'add_paragraph'):
                error_para = parent.add_paragraph(DEFAULT_EMPTY_VALUE_TEXT)
            else:
                error_para = doc.add_paragraph(DEFAULT_EMPTY_VALUE_TEXT)
            # Apply placeholder style
            run = error_para.runs[0]
            run.font.bold = DEFAULT_EMPTY_VALUE_STYLE.get("bold", True)
            if DEFAULT_EMPTY_VALUE_STYLE.get("font_color"):
                run.font.color.rgb = RGBColor.from_string(DEFAULT_EMPTY_VALUE_STYLE["font_color"])
            parent.insert(index, error_para._element)
            return

        parent.insert(index, new_para._element)


def _apply_image_wrapping(run, style_kwargs):
    """
    Apply text wrapping and positioning properties to an image.
    
    Args:
        run: The run containing the image
        style_kwargs: Style keyword arguments containing wrapping properties
    """
    # Get the shape (image) from the run
    if not run._element.findall('.//pic:pic', {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}):
        return  # No image found in run
    
    # Get the shape element
    shape = run._element.findall('.//wp:anchor', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
    if not shape:
        # If no anchor found, the image is inline - we need to convert it to floating
        # This is a complex operation that requires creating a new anchor element
        # For now, we'll only apply wrapping to images that are already floating
        return
    
    shape = shape[0]
    
    # Apply text wrapping
    wrap_text = style_kwargs.get("wrap_text")
    if wrap_text:
        _set_text_wrapping(shape, wrap_text)
    
    # Apply positioning
    horizontal_align = style_kwargs.get("horizontal_align")
    vertical_align = style_kwargs.get("vertical_align")
    if horizontal_align or vertical_align:
        _set_image_positioning(shape, horizontal_align, vertical_align)
    
    # Apply distance from text
    distance = style_kwargs.get("distance_from_text")
    if distance:
        _set_distance_from_text(shape, distance)


def _set_text_wrapping(shape, wrap_mode):
    """
    Set the text wrapping mode for an image.
    
    Args:
        shape: The shape element
        wrap_mode: The wrapping mode string
    """
    # Define the wrapping mode mappings
    wrap_map = {
        "inline": "inline",
        "square": "square",
        "tight": "tight",
        "through": "through",
        "top_and_bottom": "topAndBottom",
        "behind": "behind",
        "in_front": "inFront"
    }
    
    if wrap_mode in wrap_map:
        # Find or create the wrap element
        wrap_elem = shape.find('.//wp:wrap', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
        if wrap_elem is None:
            # Create wrap element if it doesn't exist
            from docx.oxml import parse_xml
            wrap_xml = f'<wp:wrap xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" type="{wrap_map[wrap_mode]}"/>'
            wrap_elem = parse_xml(wrap_xml)
            shape.append(wrap_elem)
        else:
            # Update existing wrap element
            wrap_elem.set('type', wrap_map[wrap_mode])


def _set_image_positioning(shape, horizontal_align, vertical_align):
    """
    Set the horizontal and vertical positioning of an image.
    
    Args:
        shape: The shape element
        horizontal_align: Horizontal alignment ("left", "center", "right")
        vertical_align: Vertical alignment ("top", "middle", "bottom")
    """
    # Define alignment mappings
    horizontal_map = {
        "left": "left",
        "center": "center", 
        "right": "right"
    }
    
    vertical_map = {
        "top": "top",
        "middle": "middle",
        "bottom": "bottom"
    }
    
    # Apply horizontal alignment
    if horizontal_align and horizontal_align in horizontal_map:
        pos_h = shape.find('.//wp:positionH', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
        if pos_h is not None:
            align_elem = pos_h.find('.//wp:align', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
            if align_elem is not None:
                align_elem.text = horizontal_map[horizontal_align]
    
    # Apply vertical alignment
    if vertical_align and vertical_align in vertical_map:
        pos_v = shape.find('.//wp:positionV', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
        if pos_v is not None:
            align_elem = pos_v.find('.//wp:align', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
            if align_elem is not None:
                align_elem.text = vertical_map[vertical_align]


def _set_distance_from_text(shape, distance):
    """
    Set the distance from text for an image.
    
    Args:
        shape: The shape element
        distance: Distance string (e.g., "0.1in", "10px")
    """
    # Parse the distance
    distance_in = _parse_measurement(distance)
    if distance_in is None:
        return
    
    # Convert to EMUs (Excel Metric Units - 1 inch = 914400 EMUs)
    distance_emu = int(distance_in * 914400)
    
    # Apply to wrap margins
    wrap_elem = shape.find('.//wp:wrap', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
    if wrap_elem is not None:
        # Set all margins to the specified distance
        for margin in ['left', 'right', 'top', 'bottom']:
            margin_elem = wrap_elem.find(f'.//wp:{margin}', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
            if margin_elem is not None:
                margin_elem.text = str(distance_emu)


def _parse_measurement(value):
    """
    Parse measurement strings and convert to inches.
    
    Accepts strings like '4in' or '300px' and returns inches as float.
    Supports only inches and pixels for now.
    
    Args:
        value: Measurement string (e.g., "4in", "300px") or None
        
    Returns:
        float: Measurement in inches, or None if parsing fails
    """
    if not value or not isinstance(value, str):
        return None

    value = value.strip().lower()

    if value.endswith("in"):
        try:
            return float(value.replace("in", ""))
        except ValueError:
            return None
    elif value.endswith("px"):
        try:
            px = float(value.replace("px", ""))
            return px / 96.0  # assuming 96 dpi standard
        except ValueError:
            return None
    else:
        return None
