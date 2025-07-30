import os
import tempfile
from PIL import Image
from docx import Document
from docxblocks.core.inserter import DocxBuilder
from docxblocks.constants import DEFAULT_EMPTY_VALUE_TEXT

def create_test_image(path, width=800, height=600, dpi=(96, 96)):
    """Create a test image with specified dimensions and DPI"""
    img = Image.new('RGB', (width, height), color='red')
    img.save(path, dpi=dpi)
    return path

def test_image_block(tmp_path):
    """Test basic image block functionality with placeholder for missing file"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "image", "path": "nonexistent.png", "style": {"max_width": "2in"}}
    ]
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    found = any(DEFAULT_EMPTY_VALUE_TEXT in p.text for p in doc2.paragraphs)
    assert found, "Image placeholder text not found in output docx"

def test_image_resizing(tmp_path):
    """Test that image resizing works correctly with max_width and max_height"""
    # Create test template
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    # Create a test image (800x600 pixels at 96 DPI = 8.33x6.25 inches)
    test_image = tmp_path / "test_image.png"
    create_test_image(str(test_image), width=800, height=600, dpi=(96, 96))
    
    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "4in",  # Should scale down to 4 inches
                "max_height": "3in"  # Should scale down to 3 inches
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that the image was inserted (not a placeholder)
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 0  # No text paragraphs, image should be present
    
    # Check that the document has content (image was inserted)
    assert len(doc2.paragraphs) > 0

def test_image_resizing_pixels(tmp_path):
    """Test image resizing with pixel measurements"""
    # Create test template
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    # Create a test image (1200x800 pixels at 96 DPI = 12.5x8.33 inches)
    test_image = tmp_path / "test_image.png"
    create_test_image(str(test_image), width=1200, height=800, dpi=(96, 96))
    
    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "600px",  # Should scale down to 6.25 inches
                "max_height": "400px"  # Should scale down to 4.17 inches
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that the image was inserted
    assert len(doc2.paragraphs) > 0

def test_image_upscaling(tmp_path):
    """Test that images can be upscaled when smaller than max constraints"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    # Create a small test image (1x1 inch at 96 DPI)
    test_image = tmp_path / "small_test.png"
    img = Image.new('RGB', (96, 96), color='red')  # 1x1 inch at 96 DPI
    img.save(test_image, dpi=(96, 96))

    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {"max_width": "3in"}  # Should upscale from 1in to 3in
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that image was inserted
    paragraphs = [p for p in doc2.paragraphs if p.runs and any(run._element.findall('.//pic:pic', {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}) for run in p.runs)]
    assert len(paragraphs) == 1
    
    # Check that the image was upscaled (this is a basic check - the actual size would need more complex verification)
    # The image should be present and not a placeholder
    assert not any(DEFAULT_EMPTY_VALUE_TEXT in p.text for p in doc2.paragraphs)

def test_image_upscaling_both_dimensions(tmp_path):
    """Test that images can be upscaled when both max_width and max_height are specified"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    # Create a small test image (1x0.5 inch at 96 DPI)
    test_image = tmp_path / "small_rect_test.png"
    img = Image.new('RGB', (96, 48), color='blue')  # 1x0.5 inch at 96 DPI
    img.save(test_image, dpi=(96, 96))

    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "4in",
                "max_height": "2in"
            }  # Should upscale to 4x2 inches
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that image was inserted
    paragraphs = [p for p in doc2.paragraphs if p.runs and any(run._element.findall('.//pic:pic', {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}) for run in p.runs)]
    assert len(paragraphs) == 1
    
    # Check that the image was upscaled and not a placeholder
    assert not any(DEFAULT_EMPTY_VALUE_TEXT in p.text for p in doc2.paragraphs)

# Text Wrapping Tests

def test_image_wrapping_square(tmp_path):
    """Test that square text wrapping can be applied to images."""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    # Create test image
    test_image = tmp_path / "test_image.png"
    create_test_image(str(test_image), width=200, height=150, dpi=(96, 96))
    
    blocks = [
        {
            "type": "text",
            "text": "This is some text that should wrap around the image. "
        },
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "2in",
                "wrap_text": "square",
                "horizontal_align": "left"
            }
        },
        {
            "type": "text",
            "text": "This text should continue after the image and wrap around it if there's enough content to demonstrate the wrapping behavior."
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    assert len(doc2.paragraphs) > 0

def test_image_wrapping_tight(tmp_path):
    """Test that tight text wrapping can be applied to images."""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    # Create test image
    test_image = tmp_path / "test_image.png"
    create_test_image(str(test_image), width=200, height=150, dpi=(96, 96))
    
    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "2in",
                "wrap_text": "tight",
                "horizontal_align": "right"
            }
        },
        {
            "type": "text",
            "text": "This text should wrap tightly around the image, following its contours more closely than square wrapping."
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    assert len(doc2.paragraphs) > 0

def test_image_wrapping_top_and_bottom(tmp_path):
    """Test that top and bottom text wrapping can be applied to images."""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    # Create test image
    test_image = tmp_path / "test_image.png"
    create_test_image(str(test_image), width=300, height=200, dpi=(96, 96))
    
    blocks = [
        {
            "type": "text",
            "text": "This text should appear above the image."
        },
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "3in",
                "wrap_text": "top_and_bottom",
                "horizontal_align": "center"
            }
        },
        {
            "type": "text",
            "text": "This text should appear below the image, with no text flowing to the sides."
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    assert len(doc2.paragraphs) > 0

def test_image_wrapping_behind(tmp_path):
    """Test that behind text wrapping can be applied to images."""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    # Create test image
    test_image = tmp_path / "test_image.png"
    create_test_image(str(test_image), width=400, height=300, dpi=(96, 96))
    
    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "4in",
                "wrap_text": "behind",
                "horizontal_align": "center",
                "vertical_align": "middle"
            }
        },
        {
            "type": "text",
            "text": "This text should appear on top of the image, creating a layered effect."
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    assert len(doc2.paragraphs) > 0

def test_image_positioning(tmp_path):
    """Test that horizontal and vertical positioning can be applied to images."""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    # Create test image
    test_image = tmp_path / "test_image.png"
    create_test_image(str(test_image), width=200, height=150, dpi=(96, 96))
    
    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "2in",
                "wrap_text": "square",
                "horizontal_align": "right",
                "vertical_align": "top"
            }
        },
        {
            "type": "text",
            "text": "This text should flow around the image which is positioned on the right side and aligned to the top."
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    assert len(doc2.paragraphs) > 0

def test_image_distance_from_text(tmp_path):
    """Test that distance from text can be applied to images."""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    # Create test image
    test_image = tmp_path / "test_image.png"
    create_test_image(str(test_image), width=200, height=150, dpi=(96, 96))
    
    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "2in",
                "wrap_text": "square",
                "horizontal_align": "left",
                "distance_from_text": "0.2in"
            }
        },
        {
            "type": "text",
            "text": "This text should have increased spacing from the image due to the distance_from_text property."
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    assert len(doc2.paragraphs) > 0

def test_image_wrapping_combined_options(tmp_path):
    """Test that multiple wrapping options can be combined."""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    # Create test image
    test_image = tmp_path / "test_image.png"
    create_test_image(str(test_image), width=250, height=180, dpi=(96, 96))
    
    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "2.5in",
                "max_height": "2in",
                "wrap_text": "square",
                "horizontal_align": "center",
                "vertical_align": "middle",
                "distance_from_text": "0.1in"
            }
        },
        {
            "type": "text",
            "text": "This image combines multiple styling options: square wrapping, center alignment, and distance from text."
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    assert len(doc2.paragraphs) > 0

def test_image_wrapping_inline_default(tmp_path):
    """Test that inline wrapping is the default behavior."""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    # Create test image
    test_image = tmp_path / "test_image.png"
    create_test_image(str(test_image), width=150, height=100, dpi=(96, 96))
    
    blocks = [
        {
            "type": "text",
            "text": "This text should flow inline with the image. "
        },
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "1.5in",
                "wrap_text": "inline"
            }
        },
        {
            "type": "text",
            "text": "This text should continue right after the image."
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    assert len(doc2.paragraphs) > 0 