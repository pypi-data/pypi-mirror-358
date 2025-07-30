import os
from docx import Document
from docxblocks.core.inserter import DocxBuilder
from PIL import Image

def test_table_block(tmp_path):
    """Test basic table block functionality"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Name", "Status", "Notes"],
                "rows": [
                    ["Service A", "OK", "All systems operational"],
                    ["Service B", "Warning", "High memory usage"],
                    ["Service C", "Error", "Connection timeout"]
                ]
            },
            "style": {
                "header_styles": {"bold": True, "bg_color": "f2f2f2"},
                "column_widths": [0.3, 0.2, 0.5]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1
    
    table = tables[0]
    assert len(table.rows) == 4  # 1 header + 3 data rows
    assert len(table.columns) == 3

def test_table_block_with_integers(tmp_path):
    """Test table block with integer values, each value is its own paragraph"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["ID", "Count", "Percentage"],
                "rows": [
                    [1, 100, 25],
                    [2, 200, 50],
                    [3, 100, 25]
                ]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Check that table was created with integer values
    tables = doc2.tables
    assert len(tables) == 1

    table = tables[0]
    # Check that integer values are properly converted to strings, each as its own paragraph
    assert table.cell(1, 0).paragraphs[0].text.strip() == "1"
    assert table.cell(2, 0).paragraphs[0].text.strip() == "2"
    assert table.cell(3, 0).paragraphs[0].text.strip() == "3"

def test_table_cell_with_single_newlines(tmp_path):
    """Test that single \n in table cells creates new paragraphs"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Name", "Description"],
                "rows": [
                    ["Item 1", "Line 1\nLine 2\nLine 3"],
                    ["Item 2", "Single line"]
                ]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1

    table = tables[0]

    # Check first row - should have 3 paragraphs for the description cell
    first_cell = table.cell(1, 1)  # Row 1, Column 1 (description)
    paragraphs = first_cell.paragraphs
    assert len(paragraphs) == 3
    assert paragraphs[0].text == "Line 1"
    assert paragraphs[1].text == "Line 2"
    assert paragraphs[2].text == "Line 3"

def test_table_cell_with_double_newlines(tmp_path):
    """Test that \n\n in table cells creates new paragraphs with blank lines"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Name", "Description"],
                "rows": [
                    ["Item 1", "First paragraph.\n\nSecond paragraph with blank line."],
                    ["Item 2", "Single line\nAnother line\n\nNew paragraph."]
                ]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1

    table = tables[0]

    # Check first row - should have 3 paragraphs for the description cell (including blank)
    first_cell = table.cell(1, 1)  # Row 1, Column 1 (description)
    paragraphs = first_cell.paragraphs
    assert len(paragraphs) == 3
    assert paragraphs[0].text == "First paragraph."
    assert paragraphs[1].text == ""
    assert paragraphs[2].text == "Second paragraph with blank line."
    # Second row, second cell
    second_cell = table.cell(2, 1)
    paragraphs2 = second_cell.paragraphs
    assert len(paragraphs2) == 4
    assert paragraphs2[0].text == "Single line"
    assert paragraphs2[1].text == "Another line"
    assert paragraphs2[2].text == ""
    assert paragraphs2[3].text == "New paragraph."

def test_table_header_with_double_newlines(tmp_path):
    """Test that \n\n in table headers creates new paragraphs with blank lines"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Name", "Description\n\nDetails"],
                "rows": [
                    ["Item 1", "Some description"],
                    ["Item 2", "Another description"]
                ]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1

    table = tables[0]

    # Check header cell with double newlines
    header_cell = table.cell(0, 1)  # Header row, Column 1
    paragraphs = header_cell.paragraphs
    assert len(paragraphs) == 3
    assert paragraphs[0].text == "Description"
    assert paragraphs[1].text == ""
    assert paragraphs[2].text == "Details"

def test_table_with_rich_cells(tmp_path):
    """Test that table cells can contain rich content (text, images, bullets, headings)"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    # Create a test image
    test_image = tmp_path / "test_image.png"
    img = Image.new('RGB', (100, 50), color='red')
    img.save(test_image, dpi=(96, 96))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Product", "Description", "Features"],
                "rows": [
                    [
                        "Product A",
                        [
                            {"type": "text", "text": "High-quality product", "style": {"bold": True}},
                            {"type": "text", "text": "Available in multiple colors", "style": {"italic": True}}
                        ],
                        [
                            {"type": "bullets", "items": ["Feature 1", "Feature 2", "Feature 3"]}
                        ]
                    ],
                    [
                        "Product B",
                        "Simple text description",
                        [
                            {"type": "image", "path": str(test_image), "style": {"max_width": "1in"}}
                        ]
                    ]
                ]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1
    
    table = tables[0]
    assert len(table.rows) == 3  # 1 header + 2 data rows
    assert len(table.columns) == 3

def test_table_with_mixed_cell_content(tmp_path):
    """Test that tables can have a mix of plain text and rich content cells"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Name", "Details", "Status"],
                "rows": [
                    [
                        "Item 1",
                        [
                            {"type": "heading", "text": "Item 1", "level": 3},
                            {"type": "text", "text": "Description with rich content."}
                        ],
                        "Active"
                    ],
                    [
                        "Item 2",
                        "Simple text description",
                        "Inactive"
                    ]
                ]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1
    
    table = tables[0]
    assert len(table.rows) == 3  # 1 header + 2 data rows
    
    # Check that mixed content works correctly
    # First row should have rich content in the details cell
    details_cell = table.cell(1, 1)  # Row 1, Column 1 (details)
    assert len(details_cell.paragraphs) > 1  # Should have multiple paragraphs from rich content
    
    # Second row should have simple text
    details_cell2 = table.cell(2, 1)  # Row 2, Column 1 (details)
    assert len(details_cell2.paragraphs) == 1  # Should have single paragraph for plain text

def test_table_rich_cells_backward_compatibility(tmp_path):
    """Test that existing plain text cells still work correctly with rich cell support"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Name", "Value"],
                "rows": [
                    ["Item 1", "100"],
                    ["Item 2", "200"],
                    ["Item 3", "300"]
                ]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1
    
    table = tables[0]
    assert len(table.rows) == 4  # 1 header + 3 data rows
    
    # Check that cells contain the expected text (backward compatibility)
    assert table.cell(1, 0).text.strip() == "Item 1"
    assert table.cell(1, 1).text.strip() == "100"
    assert table.cell(2, 0).text.strip() == "Item 2"
    assert table.cell(2, 1).text.strip() == "200" 