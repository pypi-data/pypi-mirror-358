"""
Test newline behavior with \n and \n\n
"""

import pytest
from docx import Document
from docxblocks import DocxBuilder
from docxblocks.schema.blocks import TextBlock
from docxblocks.builders.text import TextBuilder


class TestNewlineWithNewParagraph:
    """Test newline behavior with \n and \n\n only"""
    
    def test_new_paragraph_with_double_newline_start(self, tmp_path):
        """Test that \n\n at start creates an empty line before the paragraph"""
        template = tmp_path / "template.docx"
        output = tmp_path / "output.docx"
        doc = Document()
        doc.add_paragraph("{{main}}")
        doc.save(str(template))
        
        blocks = [
            {"type": "text", "text": "First paragraph"},
            {"type": "text", "text": "\n\nPlease feel free to contact me should you have further queries. I can be contacted on "},
            {"type": "text", "text": "123-456-7890"},
            {"type": "text", "text": " or "},
            {"type": "text", "text": "test@example.com"},
        ]
        
        builder = DocxBuilder(str(template))
        builder.insert("{{main}}", blocks)
        builder.save(str(output))
        
        doc2 = Document(str(output))
        paragraphs = [p.text for p in doc2.paragraphs]
        assert paragraphs[0].strip() == "First paragraph"
        # Find the first non-empty paragraph after the blank(s)
        non_empty = [p for p in paragraphs if p.strip()]
        # The first is 'First paragraph', the second is the grouped text
        assert "Please feel free to contact me should you have further queries." in non_empty[1]
        assert "123-456-7890 or test@example.com" in non_empty[1]
    
    def test_new_paragraph_with_single_newline(self, tmp_path):
        """Test that \n at start creates a new paragraph"""
        template = tmp_path / "template.docx"
        output = tmp_path / "output.docx"
        doc = Document()
        doc.add_paragraph("{{main}}")
        doc.save(str(template))
        
        blocks = [
            {"type": "text", "text": "First paragraph"},
            {"type": "text", "text": "\nPlease contact me on "},
            {"type": "text", "text": "123-456-7890"},
        ]
        
        builder = DocxBuilder(str(template))
        builder.insert("{{main}}", blocks)
        builder.save(str(output))
        
        doc2 = Document(str(output))
        paragraphs = [p.text for p in doc2.paragraphs]
        # Find the first non-empty paragraph after the blank(s)
        non_empty = [p for p in paragraphs if p.strip()]
        assert non_empty[0].strip() == "First paragraph"
        assert "Please contact me on 123-456-7890" in non_empty[1]
    
    def test_inline_text_grouping_after_new_paragraph(self, tmp_path):
        """Test that inline text grouping works correctly after a paragraph break"""
        template = tmp_path / "template.docx"
        output = tmp_path / "output.docx"
        doc = Document()
        doc.add_paragraph("{{main}}")
        doc.save(str(template))
        
        blocks = [
            {"type": "text", "text": "First paragraph"},
            {"type": "text", "text": "\nSecond paragraph"},
            {"type": "text", "text": " continues inline"},
            {"type": "text", "text": " and more inline"},
        ]
        
        builder = DocxBuilder(str(template))
        builder.insert("{{main}}", blocks)
        builder.save(str(output))
        
        doc2 = Document(str(output))
        paragraphs = [p.text for p in doc2.paragraphs]
        # Find the first non-empty paragraph after the blank(s)
        non_empty = [p for p in paragraphs if p.strip()]
        assert non_empty[0].strip() == "First paragraph"
        assert non_empty[1].startswith("Second paragraph continues inline and more inline") 