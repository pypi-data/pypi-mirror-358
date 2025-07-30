"""
Tests for rich content in table cells.
"""

import pytest
import tempfile
import os
from PIL import Image
from docx import Document
from docxblocks.core.inserter import DocxBuilder


def create_test_image(path, width=100, height=50, color='red'):
    """Create a test image with specified dimensions"""
    img = Image.new('RGB', (width, height), color=color)
    img.save(path, dpi=(96, 96))
    return path


def test_rich_cells_with_text_and_images():
    """Test that cells can contain rich content including text and images."""
    with tempfile.TemporaryDirectory() as temp_dir:
        template_path = os.path.join(temp_dir, "test_template.docx")
        output_path = os.path.join(temp_dir, "test_output.docx")
        
        # Create test images
        logo_path = os.path.join(temp_dir, "logo.png")
        create_test_image(logo_path, width=100, height=50, color='blue')
        
        doc = Document()
        doc.add_paragraph("{{main}}")
        doc.save(template_path)
        
        blocks = [
            {
                "type": "table",
                "content": {
                    "headers": ["Product", "Description", "Image"],
                    "rows": [
                        [
                            "Product A",
                            [
                                {"type": "text", "text": "High-quality product", "style": {"bold": True}},
                                {"type": "text", "text": "Available in multiple colors", "style": {"italic": True}}
                            ],
                            [
                                {"type": "image", "path": logo_path, "style": {"max_width": "1in"}}
                            ]
                        ],
                        [
                            "Product B",
                            "Simple text description",
                            "No image"
                        ]
                    ]
                }
            }
        ]
        
        builder = DocxBuilder(template_path)
        builder.insert("{{main}}", blocks)
        builder.save(output_path)
        
        # Verify the document was created successfully
        assert os.path.exists(output_path)
        result_doc = Document(output_path)
        
        # Check that table was created
        tables = result_doc.tables
        assert len(tables) == 1
        
        table = tables[0]
        assert len(table.rows) == 3  # 1 header + 2 data rows
        assert len(table.columns) == 3


def test_rich_cells_with_bullets():
    """Test that cells can contain bullet lists."""
    with tempfile.TemporaryDirectory() as temp_dir:
        template_path = os.path.join(temp_dir, "test_template.docx")
        output_path = os.path.join(temp_dir, "test_output.docx")
        
        doc = Document()
        doc.add_paragraph("{{main}}")
        doc.save(template_path)
        
        blocks = [
            {
                "type": "table",
                "content": {
                    "headers": ["Category", "Features"],
                    "rows": [
                        [
                            "Software",
                            [
                                {"type": "bullets", "items": ["Easy to use", "Fast performance", "Secure"]}
                            ]
                        ],
                        [
                            "Hardware",
                            [
                                {"type": "bullets", "items": ["Durable", "Energy efficient"]}
                            ]
                        ]
                    ]
                }
            }
        ]
        
        builder = DocxBuilder(template_path)
        builder.insert("{{main}}", blocks)
        builder.save(output_path)
        
        # Verify the document was created successfully
        assert os.path.exists(output_path)
        result_doc = Document(output_path)
        
        # Check that table was created
        tables = result_doc.tables
        assert len(tables) == 1
        
        table = tables[0]
        assert len(table.rows) == 3  # 1 header + 2 data rows


def test_rich_cells_with_headings():
    """Test that cells can contain headings."""
    with tempfile.TemporaryDirectory() as temp_dir:
        template_path = os.path.join(temp_dir, "test_template.docx")
        output_path = os.path.join(temp_dir, "test_output.docx")
        
        doc = Document()
        doc.add_paragraph("{{main}}")
        doc.save(template_path)
        
        blocks = [
            {
                "type": "table",
                "content": {
                    "headers": ["Section", "Content"],
                    "rows": [
                        [
                            "Introduction",
                            [
                                {"type": "heading", "text": "Welcome", "level": 2},
                                {"type": "text", "text": "This is the introduction section."}
                            ]
                        ],
                        [
                            "Conclusion",
                            [
                                {"type": "heading", "text": "Summary", "level": 3},
                                {"type": "text", "text": "This concludes our document."}
                            ]
                        ]
                    ]
                }
            }
        ]
        
        builder = DocxBuilder(template_path)
        builder.insert("{{main}}", blocks)
        builder.save(output_path)
        
        # Verify the document was created successfully
        assert os.path.exists(output_path)
        result_doc = Document(output_path)
        
        # Check that table was created
        tables = result_doc.tables
        assert len(tables) == 1


def test_mixed_content_cells():
    """Test that cells can contain a mix of different content types."""
    with tempfile.TemporaryDirectory() as temp_dir:
        template_path = os.path.join(temp_dir, "test_template.docx")
        output_path = os.path.join(temp_dir, "test_output.docx")
        
        # Create test image
        logo_path = os.path.join(temp_dir, "logo.png")
        create_test_image(logo_path, width=80, height=40, color='green')
        
        doc = Document()
        doc.add_paragraph("{{main}}")
        doc.save(template_path)
        
        blocks = [
            {
                "type": "table",
                "content": {
                    "headers": ["Item", "Details"],
                    "rows": [
                        [
                            "Product X",
                            [
                                {"type": "heading", "text": "Product X", "level": 3},
                                {"type": "text", "text": "Description:", "style": {"bold": True}},
                                {"type": "text", "text": "High-quality product with excellent features."},
                                {"type": "bullets", "items": ["Feature 1", "Feature 2", "Feature 3"]},
                                {"type": "image", "path": logo_path, "style": {"max_width": "0.8in"}}
                            ]
                        ]
                    ]
                }
            }
        ]
        
        builder = DocxBuilder(template_path)
        builder.insert("{{main}}", blocks)
        builder.save(output_path)
        
        # Verify the document was created successfully
        assert os.path.exists(output_path)
        result_doc = Document(output_path)
        
        # Check that table was created
        tables = result_doc.tables
        assert len(tables) == 1


def test_backward_compatibility():
    """Test that existing plain text cells still work correctly."""
    with tempfile.TemporaryDirectory() as temp_dir:
        template_path = os.path.join(temp_dir, "test_template.docx")
        output_path = os.path.join(temp_dir, "test_output.docx")
        
        doc = Document()
        doc.add_paragraph("{{main}}")
        doc.save(template_path)
        
        blocks = [
            {
                "type": "table",
                "content": {
                    "headers": ["Name", "Value"],
                    "rows": [
                        ["Item 1", "100"],
                        ["Item 2", "200"],
                        ["Item 3", "300"]
                    ]
                }
            }
        ]
        
        builder = DocxBuilder(template_path)
        builder.insert("{{main}}", blocks)
        builder.save(output_path)
        
        # Verify the document was created successfully
        assert os.path.exists(output_path)
        result_doc = Document(output_path)
        
        # Check that table was created
        tables = result_doc.tables
        assert len(tables) == 1
        
        table = tables[0]
        assert len(table.rows) == 4  # 1 header + 3 data rows
        
        # Check that cells contain the expected text
        assert table.cell(1, 0).text.strip() == "Item 1"
        assert table.cell(1, 1).text.strip() == "100"


if __name__ == "__main__":
    pytest.main([__file__]) 