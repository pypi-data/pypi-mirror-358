"""
Tests for header and footer functionality.
"""

import pytest
import tempfile
import os
from docx import Document
from docxblocks.core.inserter import DocxBuilder
from docxblocks.schema.blocks import HeaderBlock, FooterBlock
from docxblocks.builders.header_footer import HeaderFooterBuilder
from PIL import Image


def test_header_block_validation():
    """Test that HeaderBlock validates correctly."""
    valid_header = {
        "type": "header",
        "content": [
            {
                "type": "text",
                "text": "Test Header"
            }
        ],
        "apply_to": "all"
    }
    
    header_block = HeaderBlock.model_validate(valid_header)
    assert header_block.type == "header"
    assert header_block.apply_to == "all"
    assert len(header_block.content) == 1


def test_footer_block_validation():
    """Test that FooterBlock validates correctly."""
    valid_footer = {
        "type": "footer",
        "content": [
            {
                "type": "text", 
                "text": "Test Footer"
            }
        ],
        "apply_to": "first"
    }
    
    footer_block = FooterBlock.model_validate(valid_footer)
    assert footer_block.type == "footer"
    assert footer_block.apply_to == "first"
    assert len(footer_block.content) == 1


def test_header_footer_builder_initialization():
    """Test HeaderFooterBuilder initialization."""
    doc = Document()
    builder = HeaderFooterBuilder(doc)
    assert builder.doc == doc


def test_basic_header_footer():
    """Test basic header and footer functionality."""
    with tempfile.TemporaryDirectory() as temp_dir:
        # Create template
        template_path = os.path.join(temp_dir, "test_template.docx")
        output_path = os.path.join(temp_dir, "test_output.docx")
        
        doc = Document()
        doc.add_paragraph("{{content}}")
        doc.save(template_path)
        
        # Define blocks with header and footer
        blocks = [
            {
                "type": "header",
                "apply_to": "all",
                "content": [
                    {
                        "type": "text",
                        "text": "Test Header",
                        "style": {"align": "center"}
                    }
                ]
            },
            {
                "type": "footer", 
                "apply_to": "all",
                "content": [
                    {
                        "type": "text",
                        "text": "Test Footer",
                        "style": {"align": "center"}
                    }
                ]
            },
            {
                "type": "text",
                "text": "Main content"
            }
        ]
        
        # Build document
        builder = DocxBuilder(template_path)
        builder.insert("{{content}}", blocks)
        builder.save(output_path)
        
        # Verify the document was created
        assert os.path.exists(output_path)
        
        # Verify the document can be opened and has headers/footers
        result_doc = Document(output_path)
        section = result_doc.sections[0]
        
        # Check that header and footer are not linked to previous (meaning they have content)
        assert not section.header.is_linked_to_previous
        assert not section.footer.is_linked_to_previous


def test_first_page_header_footer():
    """Test different header/footer for first page."""
    with tempfile.TemporaryDirectory() as temp_dir:
        template_path = os.path.join(temp_dir, "test_template.docx")
        output_path = os.path.join(temp_dir, "test_output.docx")
        
        doc = Document()
        doc.add_paragraph("{{content}}")
        doc.save(template_path)
        
        blocks = [
            {
                "type": "header",
                "apply_to": "first",
                "content": [
                    {
                        "type": "text",
                        "text": "First Page Header"
                    }
                ]
            },
            {
                "type": "footer",
                "apply_to": "first", 
                "content": [
                    {
                        "type": "text",
                        "text": "First Page Footer"
                    }
                ]
            },
            {
                "type": "text",
                "text": "Page 1 content"
            },
            {
                "type": "page_break"
            },
            {
                "type": "text", 
                "text": "Page 2 content"
            }
        ]
        
        builder = DocxBuilder(template_path)
        builder.insert("{{content}}", blocks)
        builder.save(output_path)
        
        # Verify document settings
        result_doc = Document(output_path)
        section = result_doc.sections[0]
        
        # Check that different first page header/footer is enabled
        assert section.different_first_page_header_footer
        assert not section.first_page_header.is_linked_to_previous
        assert not section.first_page_footer.is_linked_to_previous


def test_odd_even_header_footer():
    """Test different headers/footers for odd and even pages."""
    with tempfile.TemporaryDirectory() as temp_dir:
        template_path = os.path.join(temp_dir, "test_template.docx")
        output_path = os.path.join(temp_dir, "test_output.docx")
        
        doc = Document()
        doc.add_paragraph("{{content}}")
        doc.save(template_path)
        
        blocks = [
            {
                "type": "header",
                "apply_to": "odd",
                "content": [
                    {
                        "type": "text",
                        "text": "Odd Page Header"
                    }
                ]
            },
            {
                "type": "header",
                "apply_to": "even",
                "content": [
                    {
                        "type": "text", 
                        "text": "Even Page Header"
                    }
                ]
            },
            {
                "type": "text",
                "text": "Main content"
            }
        ]
        
        builder = DocxBuilder(template_path)
        builder.insert("{{content}}", blocks)
        builder.save(output_path)
        
        # Verify document settings
        result_doc = Document(output_path)
        
        # Check that odd/even headers are enabled
        assert result_doc.settings.odd_and_even_pages_header_footer
        
        section = result_doc.sections[0]
        assert not section.header.is_linked_to_previous  # Odd page header
        assert not section.even_page_header.is_linked_to_previous


def test_header_footer_with_complex_content():
    """Test headers and footers with complex content including tables and images."""
    with tempfile.TemporaryDirectory() as temp_dir:
        template_path = os.path.join(temp_dir, "test_template.docx")
        output_path = os.path.join(temp_dir, "test_output.docx")
        
        doc = Document()
        doc.add_paragraph("{{content}}")
        doc.save(template_path)
        
        blocks = [
            {
                "type": "header",
                "apply_to": "all",
                "content": [
                    {
                        "type": "table",
                        "content": {
                            "headers": ["Company", "Document", "Page"],
                            "rows": [["ACME Corp", "Report", "{{page}}"]]
                        }
                    }
                ]
            },
            {
                "type": "footer",
                "apply_to": "all",
                "content": [
                    {
                        "type": "text",
                        "text": "Generated on: ",
                        "style": {"bold": True}
                    },
                    {
                        "type": "text", 
                        "text": "{{date}}"
                    }
                ]
            },
            {
                "type": "text",
                "text": "Document content goes here."
            }
        ]
        
        builder = DocxBuilder(template_path)
        builder.insert("{{content}}", blocks)
        builder.save(output_path)
        
        # Verify the document was created successfully
        assert os.path.exists(output_path)
        result_doc = Document(output_path)
        section = result_doc.sections[0]
        
        # Verify headers and footers have content
        assert not section.header.is_linked_to_previous
        assert not section.footer.is_linked_to_previous


def test_header_footer_schema_validation():
    """Test that header and footer blocks validate correctly."""
    
    # Test valid header block
    valid_header = {
        "type": "header",
        "apply_to": "all",
        "content": [
            {
                "type": "text",
                "text": "Header text"
            }
        ]
    }
    
    header_block = HeaderBlock.model_validate(valid_header)
    assert header_block.type == "header"
    assert header_block.apply_to == "all"
    assert len(header_block.content) == 1
    
    # Test valid footer block  
    valid_footer = {
        "type": "footer",
        "apply_to": "first", 
        "content": [
            {
                "type": "text",
                "text": "Footer text"
            }
        ]
    }
    
    footer_block = FooterBlock.model_validate(valid_footer)
    assert footer_block.type == "footer"
    assert footer_block.apply_to == "first"
    assert len(footer_block.content) == 1
    
    # Test default apply_to value
    minimal_header = {
        "type": "header",
        "content": []
    }
    
    header_block = HeaderBlock.model_validate(minimal_header)
    assert header_block.apply_to == "all"  # Default value


def test_all_except_first_header_footer():
    """Test headers and footers that apply to all pages except the first."""
    with tempfile.TemporaryDirectory() as temp_dir:
        template_path = os.path.join(temp_dir, "test_template.docx")
        output_path = os.path.join(temp_dir, "test_output.docx")
        
        doc = Document()
        doc.add_paragraph("{{content}}")
        doc.save(template_path)
        
        blocks = [
            {
                "type": "header",
                "apply_to": "all_except_first",
                "content": [
                    {
                        "type": "text",
                        "text": "Header on pages 2+"
                    }
                ]
            },
            {
                "type": "footer",
                "apply_to": "all_except_first",
                "content": [
                    {
                        "type": "text",
                        "text": "Footer on pages 2+"
                    }
                ]
            },
            {
                "type": "text",
                "text": "Page 1 content (no header/footer)"
            },
            {
                "type": "page_break"
            },
            {
                "type": "text",
                "text": "Page 2 content (with header/footer)"
            }
        ]
        
        builder = DocxBuilder(template_path)
        builder.insert("{{content}}", blocks)
        builder.save(output_path)
        
        # Verify document settings
        result_doc = Document(output_path)
        section = result_doc.sections[0]
        
        # Check that different first page is enabled
        assert section.different_first_page_header_footer
        
        # Check that first page header/footer are linked (empty)
        assert section.first_page_header.is_linked_to_previous
        assert section.first_page_footer.is_linked_to_previous
        
        # Check that default header/footer have content (for pages 2+)
        assert not section.header.is_linked_to_previous
        assert not section.footer.is_linked_to_previous


def test_header_with_image():
    """Test that images can be added to headers and display correctly."""
    with tempfile.TemporaryDirectory() as temp_dir:
        template_path = os.path.join(temp_dir, "test_template.docx")
        output_path = os.path.join(temp_dir, "test_output.docx")
        
        # Create a test image
        test_image = os.path.join(temp_dir, "test_image.png")
        img = Image.new('RGB', (200, 100), color='blue')
        img.save(test_image, dpi=(96, 96))
        
        doc = Document()
        doc.add_paragraph("{{content}}")
        doc.save(template_path)
        
        blocks = [
            {
                "type": "header",
                "apply_to": "all",
                "content": [
                    {
                        "type": "image",
                        "path": test_image,
                        "style": {"max_width": "1.5in"}
                    },
                    {
                        "type": "text",
                        "text": "Header with Image",
                        "style": {"align": "center", "bold": True}
                    }
                ]
            },
            {
                "type": "text",
                "text": "Main document content. The header should contain an image."
            }
        ]
        
        builder = DocxBuilder(template_path)
        builder.insert("{{content}}", blocks)
        builder.save(output_path)
        
        # Verify the document was created successfully
        assert os.path.exists(output_path)
        result_doc = Document(output_path)
        section = result_doc.sections[0]
        
        # Verify header has content (not linked to previous)
        assert not section.header.is_linked_to_previous
        
        # Check that the header has paragraphs (indicating content was added)
        assert len(section.header.paragraphs) > 0


def test_footer_with_image():
    """Test that images can be added to footers and display correctly."""
    with tempfile.TemporaryDirectory() as temp_dir:
        template_path = os.path.join(temp_dir, "test_template.docx")
        output_path = os.path.join(temp_dir, "test_output.docx")
        
        # Create a test image
        test_image = os.path.join(temp_dir, "test_image.png")
        img = Image.new('RGB', (200, 100), color='green')
        img.save(test_image, dpi=(96, 96))
        
        doc = Document()
        doc.add_paragraph("{{content}}")
        doc.save(template_path)
        
        blocks = [
            {
                "type": "footer",
                "apply_to": "all",
                "content": [
                    {
                        "type": "text",
                        "text": "Footer with Image",
                        "style": {"align": "center", "bold": True}
                    },
                    {
                        "type": "image",
                        "path": test_image,
                        "style": {"max_width": "1.5in"}
                    }
                ]
            },
            {
                "type": "text",
                "text": "Main document content. The footer should contain an image."
            }
        ]
        
        builder = DocxBuilder(template_path)
        builder.insert("{{content}}", blocks)
        builder.save(output_path)
        
        # Verify the document was created successfully
        assert os.path.exists(output_path)
        result_doc = Document(output_path)
        section = result_doc.sections[0]
        
        # Verify footer has content (not linked to previous)
        assert not section.footer.is_linked_to_previous
        
        # Check that the footer has paragraphs (indicating content was added)
        assert len(section.footer.paragraphs) > 0


def test_invalid_apply_to_values():
    """Test that invalid apply_to values are rejected."""
    
    invalid_header = {
        "type": "header",
        "apply_to": "invalid_value",
        "content": []
    }
    
    with pytest.raises(Exception):  # Pydantic validation error
        HeaderBlock.model_validate(invalid_header)


if __name__ == "__main__":
    pytest.main([__file__]) 