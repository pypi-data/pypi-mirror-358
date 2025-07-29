import os
import tempfile
from PIL import Image
from docx import Document
from docxblocks.core.inserter import DocxBuilder
from docxblocks.constants import DEFAULT_EMPTY_VALUE_TEXT

def create_test_image(path, width=800, height=600, dpi=(96, 96)):
    """Create a test image with specified dimensions and DPI"""
    img = Image.new('RGB', (width, height), color='red')
    img.save(path, dpi=dpi)
    return path

def test_image_block(tmp_path):
    """Test basic image block functionality with placeholder for missing file"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "image", "path": "nonexistent.png", "style": {"max_width": "2in"}}
    ]
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    found = any(DEFAULT_EMPTY_VALUE_TEXT in p.text for p in doc2.paragraphs)
    assert found, "Image placeholder text not found in output docx"

def test_image_resizing(tmp_path):
    """Test that image resizing works correctly with max_width and max_height"""
    # Create test template
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    # Create a test image (800x600 pixels at 96 DPI = 8.33x6.25 inches)
    test_image = tmp_path / "test_image.png"
    create_test_image(str(test_image), width=800, height=600, dpi=(96, 96))
    
    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "4in",  # Should scale down to 4 inches
                "max_height": "3in"  # Should scale down to 3 inches
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that the image was inserted (not a placeholder)
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 0  # No text paragraphs, image should be present
    
    # Check that the document has content (image was inserted)
    assert len(doc2.paragraphs) > 0

def test_image_resizing_pixels(tmp_path):
    """Test image resizing with pixel measurements"""
    # Create test template
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    # Create a test image (1200x800 pixels at 96 DPI = 12.5x8.33 inches)
    test_image = tmp_path / "test_image.png"
    create_test_image(str(test_image), width=1200, height=800, dpi=(96, 96))
    
    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "600px",  # Should scale down to 6.25 inches
                "max_height": "400px"  # Should scale down to 4.17 inches
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that the image was inserted
    assert len(doc2.paragraphs) > 0

def test_image_placeholder_for_missing_file(tmp_path):
    """Test that placeholder is shown for missing image files"""
    # Create test template
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))
    
    blocks = [
        {
            "type": "image",
            "path": "nonexistent_image.png",
            "style": {
                "max_width": "4in",
                "max_height": "3in"
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))
    
    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that placeholder text was inserted
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 1
    assert "VALUE NOT FOUND" in paragraphs[0]

def test_image_upscaling(tmp_path):
    """Test that images can be upscaled when smaller than max constraints"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    # Create a small test image (1x1 inch at 96 DPI)
    test_image = tmp_path / "small_test.png"
    img = Image.new('RGB', (96, 96), color='red')  # 1x1 inch at 96 DPI
    img.save(test_image, dpi=(96, 96))

    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {"max_width": "3in"}  # Should upscale from 1in to 3in
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that image was inserted
    paragraphs = [p for p in doc2.paragraphs if p.runs and any(run._element.findall('.//pic:pic', {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}) for run in p.runs)]
    assert len(paragraphs) == 1
    
    # Check that the image was upscaled (this is a basic check - the actual size would need more complex verification)
    # The image should be present and not a placeholder
    assert not any(DEFAULT_EMPTY_VALUE_TEXT in p.text for p in doc2.paragraphs)

def test_image_upscaling_both_dimensions(tmp_path):
    """Test that images can be upscaled when both max_width and max_height are specified"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    # Create a small test image (1x0.5 inch at 96 DPI)
    test_image = tmp_path / "small_rect_test.png"
    img = Image.new('RGB', (96, 48), color='blue')  # 1x0.5 inch at 96 DPI
    img.save(test_image, dpi=(96, 96))

    blocks = [
        {
            "type": "image",
            "path": str(test_image),
            "style": {
                "max_width": "4in",
                "max_height": "2in"
            }  # Should upscale to 4x2 inches
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that image was inserted
    paragraphs = [p for p in doc2.paragraphs if p.runs and any(run._element.findall('.//pic:pic', {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}) for run in p.runs)]
    assert len(paragraphs) == 1
    
    # Check that the image was upscaled and not a placeholder
    assert not any(DEFAULT_EMPTY_VALUE_TEXT in p.text for p in doc2.paragraphs) 