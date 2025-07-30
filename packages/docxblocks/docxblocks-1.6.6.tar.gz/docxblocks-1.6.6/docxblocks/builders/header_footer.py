"""
Header and Footer Builder Module

This module provides the HeaderFooterBuilder class for rendering header and footer blocks in Word documents.
It handles header and footer insertion with support for various page application rules and rich content.
"""

from docx import Document
from docxblocks.schema.blocks import HeaderBlock, FooterBlock
from docxblocks.schema.shared import TextStyle
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class HeaderFooterBuilder:
    """
    Builder class for rendering header and footer blocks in Word documents.
    
    This builder handles header and footer insertion with support for various
    page application rules (all pages, all except first, first only, odd/even).
    It supports rich content including text, images, and other block types.
    
    The builder automatically configures section properties to apply headers
    and footers according to the specified rules and handles content rendering
    using the same logic as the main document builder.
    """
    
    def __init__(self, doc):
        """
        Initialize the HeaderFooterBuilder.
        
        Args:
            doc: The python-docx Document object
        """
        self.doc = doc
    
    def build_header(self, block: HeaderBlock):
        """
        Build and render a header block.
        
        Args:
            block: A validated HeaderBlock object
        """
        # Get the section
        section = self.doc.sections[0]
        
        # Configure header based on apply_to rule
        self._configure_header_application(section, block.apply_to)
        
        # Get the appropriate header
        header = self._get_header_for_rule(section, block.apply_to)
        
        # Render content in the header
        self._configure_header(header, block.content)
    
    def build_footer(self, block: FooterBlock):
        """
        Build and render a footer block.
        
        Args:
            block: A validated FooterBlock object
        """
        # Get the section
        section = self.doc.sections[0]
        
        # Configure footer based on apply_to rule
        self._configure_footer_application(section, block.apply_to)
        
        # Get the appropriate footer
        footer = self._get_footer_for_rule(section, block.apply_to)
        
        # Render content in the footer
        self._configure_footer(footer, block.content)
    
    def _configure_header_application(self, section, apply_to):
        """
        Configure header application based on the apply_to rule.
        
        Args:
            section: The document section
            apply_to: The application rule
        """
        if apply_to in ["first", "all_except_first"]:
            section.different_first_page_header_footer = True
        elif apply_to in ["odd", "even"]:
            self.doc.settings.odd_and_even_pages_header_footer = True
    
    def _configure_footer_application(self, section, apply_to):
        """
        Configure footer application based on the apply_to rule.
        
        Args:
            section: The document section
            apply_to: The application rule
        """
        if apply_to in ["first", "all_except_first"]:
            section.different_first_page_header_footer = True
        elif apply_to in ["odd", "even"]:
            self.doc.settings.odd_and_even_pages_header_footer = True
    
    def _get_header_for_rule(self, section, apply_to):
        """
        Get the appropriate header based on the apply_to rule.
        
        Args:
            section: The document section
            apply_to: The application rule
            
        Returns:
            The appropriate header object
        """
        if apply_to == "first":
            return section.first_page_header
        elif apply_to == "odd":
            return section.header
        elif apply_to == "even":
            return section.even_page_header
        else:  # "all" or "all_except_first"
            return section.header
    
    def _get_footer_for_rule(self, section, apply_to):
        """
        Get the appropriate footer based on the apply_to rule.
        
        Args:
            section: The document section
            apply_to: The application rule
            
        Returns:
            The appropriate footer object
        """
        if apply_to == "first":
            return section.first_page_footer
        elif apply_to == "odd":
            return section.footer
        elif apply_to == "even":
            return section.even_page_footer
        else:  # "all" or "all_except_first"
            return section.footer

    def _configure_header(self, header, content_blocks):
        """
        Configure a specific header with the provided content blocks.
        
        Args:
            header: The python-docx header object (_Header)
            content_blocks: List of block dictionaries to render in the header
        """
        # Clear any existing content
        for paragraph in header.paragraphs:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        
        # Process content blocks directly using header methods
        for block in content_blocks:
            block_type = block.get('type')
            
            if block_type == 'text':
                para = header.add_paragraph()
                run = para.add_run(block.get('text', ''))
                # Apply text styling if provided
                style = block.get('style', {})
                if style.get('bold'):
                    run.bold = True
                if style.get('italic'):
                    run.italic = True
                if style.get('font_color'):
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor.from_string(style['font_color'])
                if style.get('align'):
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    align_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if style['align'] in align_map:
                        para.alignment = align_map[style['align']]
                        
            elif block_type == 'image':
                para = header.add_paragraph()
                run = para.add_run()
                
                image_path = block.get('path', '')
                style = block.get('style', {})
                
                if image_path and os.path.isfile(image_path):
                    try:
                        from PIL import Image as PILImage
                        from docx.shared import Inches
                        
                        with PILImage.open(image_path) as img:
                            width_px, height_px = img.size
                            dpi_x, dpi_y = img.info.get("dpi", (72, 72))
                            width_in = width_px / dpi_x
                            height_in = height_px / dpi_y

                            # Calculate scale factors for width and height constraints
                            scales = []
                            max_width = self._parse_measurement(style.get("max_width"))
                            max_height = self._parse_measurement(style.get("max_height"))

                            if max_width:
                                scales.append(max_width / width_in)
                            if max_height:
                                scales.append(max_height / height_in)
                            
                            # Use the minimum scale to ensure neither dimension exceeds its maximum
                            if scales:
                                scale = min(scales)
                            else:
                                scale = 1.0

                            # Use file object to ensure embedding works in headers/footers
                            with open(image_path, 'rb') as img_file:
                                run.add_picture(img_file, width=Inches(width_in * scale), height=Inches(height_in * scale))
                            
                            # Apply basic text wrapping only - avoid complex XML manipulation
                            wrap_text = style.get("wrap_text")
                            if wrap_text and wrap_text != "inline":
                                # Only apply simple wrapping, avoid complex conversion
                                self._apply_simple_wrapping(run, wrap_text)
                            
                    except Exception as e:
                        # Handle image loading errors gracefully
                        run.text = f"[Image: {os.path.basename(image_path)}]"
                else:
                    # Handle missing image
                    run.text = f"[Missing Image: {image_path}]"

    def _configure_footer(self, footer, content_blocks):
        """
        Configure a specific footer with the provided content blocks.
        
        Args:
            footer: The python-docx footer object (_Footer)
            content_blocks: List of block dictionaries to render in the footer
        """
        # Clear any existing content
        for paragraph in footer.paragraphs:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        
        # Process content blocks directly using footer methods
        for block in content_blocks:
            block_type = block.get('type')
            
            if block_type == 'text':
                para = footer.add_paragraph()
                run = para.add_run(block.get('text', ''))
                # Apply text styling if provided
                style = block.get('style', {})
                if style.get('bold'):
                    run.bold = True
                if style.get('italic'):
                    run.italic = True
                if style.get('font_color'):
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor.from_string(style['font_color'])
                if style.get('align'):
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    align_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if style['align'] in align_map:
                        para.alignment = align_map[style['align']]
                        
            elif block_type == 'image':
                para = footer.add_paragraph()
                run = para.add_run()
                
                image_path = block.get('path', '')
                style = block.get('style', {})
                
                if image_path and os.path.isfile(image_path):
                    try:
                        from PIL import Image as PILImage
                        from docx.shared import Inches
                        
                        with PILImage.open(image_path) as img:
                            width_px, height_px = img.size
                            dpi_x, dpi_y = img.info.get("dpi", (72, 72))
                            width_in = width_px / dpi_x
                            height_in = height_px / dpi_y

                            # Calculate scale factors for width and height constraints
                            scales = []
                            max_width = self._parse_measurement(style.get("max_width"))
                            max_height = self._parse_measurement(style.get("max_height"))

                            if max_width:
                                scales.append(max_width / width_in)
                            if max_height:
                                scales.append(max_height / height_in)
                            
                            # Use the minimum scale to ensure neither dimension exceeds its maximum
                            if scales:
                                scale = min(scales)
                            else:
                                scale = 1.0

                            # Use file object to ensure embedding works in headers/footers
                            with open(image_path, 'rb') as img_file:
                                run.add_picture(img_file, width=Inches(width_in * scale), height=Inches(height_in * scale))
                            
                            # Apply basic text wrapping only - avoid complex XML manipulation
                            wrap_text = style.get("wrap_text")
                            if wrap_text and wrap_text != "inline":
                                # Only apply simple wrapping, avoid complex conversion
                                self._apply_simple_wrapping(run, wrap_text)
                            
                    except Exception as e:
                        # Handle image loading errors gracefully
                        run.text = f"[Image: {os.path.basename(image_path)}]"
                else:
                    # Handle missing image
                    run.text = f"[Missing Image: {image_path}]"

    def _apply_simple_wrapping(self, run, wrap_mode):
        """
        Apply simple text wrapping without complex XML manipulation.
        
        Args:
            run: The run containing the image
            wrap_mode: The wrapping mode string
        """
        # For now, just log the wrapping request but don't apply complex XML changes
        # This prevents corruption while still allowing basic image embedding
        pass

    def _parse_measurement(self, value):
        """
        Parse measurement strings and convert to inches.
        
        Accepts strings like '4in' or '300px' and returns inches as float.
        Supports only inches and pixels for now.
        
        Args:
            value: Measurement string (e.g., "4in", "300px") or None
            
        Returns:
            float: Measurement in inches, or None if parsing fails
        """
        if not value or not isinstance(value, str):
            return None

        value = value.strip().lower()

        if value.endswith("in"):
            try:
                return float(value.replace("in", ""))
            except ValueError:
                return None
        elif value.endswith("px"):
            try:
                px = float(value.replace("px", ""))
                return px / 96.0  # assuming 96 dpi standard
            except ValueError:
                return None
        else:
            return None 