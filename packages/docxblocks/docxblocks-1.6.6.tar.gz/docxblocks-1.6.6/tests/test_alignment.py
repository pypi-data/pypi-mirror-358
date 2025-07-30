import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxblocks.core.inserter import DocxBuilder

def test_text_block_alignment(tmp_path):
    """Test text block alignment for all alignment options including justify"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "text",
            "text": "Left aligned text",
            "style": {"align": "left"}
        },
        {
            "type": "text", 
            "text": "Center aligned text",
            "style": {"align": "center"}
        },
        {
            "type": "text",
            "text": "Right aligned text", 
            "style": {"align": "right"}
        },
        {
            "type": "text",
            "text": "Justified text that should be aligned to both left and right margins for proper testing",
            "style": {"align": "justify"}
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that paragraphs exist and have correct alignment
    paragraphs = [p for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 4
    
    assert paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert paragraphs[3].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

def test_heading_alignment(tmp_path):
    """Test heading alignment for all alignment options including justify"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "heading",
            "text": "Left Heading",
            "level": 1,
            "style": {"align": "left"}
        },
        {
            "type": "heading",
            "text": "Center Heading", 
            "level": 2,
            "style": {"align": "center"}
        },
        {
            "type": "heading",
            "text": "Right Heading",
            "level": 3,
            "style": {"align": "right"}
        },
        {
            "type": "heading",
            "text": "Justified Heading",
            "level": 4,
            "style": {"align": "justify"}
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that paragraphs exist and have correct alignment
    paragraphs = [p for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 4
    
    assert paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert paragraphs[3].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

def test_bullet_alignment(tmp_path):
    """Test bullet list alignment for all alignment options including justify"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "bullets",
            "items": ["Left item 1", "Left item 2"],
            "style": {"align": "left"}
        },
        {
            "type": "bullets",
            "items": ["Center item 1", "Center item 2"],
            "style": {"align": "center"}
        },
        {
            "type": "bullets",
            "items": ["Right item 1", "Right item 2"],
            "style": {"align": "right"}
        },
        {
            "type": "bullets",
            "items": ["Justified item with longer text to demonstrate alignment", "Another justified item"],
            "style": {"align": "justify"}
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that paragraphs exist and have correct alignment
    paragraphs = [p for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 8  # 2 items per bullet list * 4 lists
    
    # Check first bullet list (left)
    assert paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.LEFT
    
    # Check second bullet list (center)
    assert paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraphs[3].alignment == WD_ALIGN_PARAGRAPH.CENTER
    
    # Check third bullet list (right)
    assert paragraphs[4].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert paragraphs[5].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    
    # Check fourth bullet list (justify)
    assert paragraphs[6].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert paragraphs[7].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

def test_table_header_alignment(tmp_path):
    """Test table header alignment for all alignment options including justify"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Left Header", "Center Header", "Right Header", "Justified Header"],
                "rows": [
                    ["Data 1", "Data 2", "Data 3", "Data 4"]
                ]
            },
            "style": {
                "header_styles": {"align": "left"}
            }
        },
        {
            "type": "table",
            "content": {
                "headers": ["Center Test"],
                "rows": [["Data"]]
            },
            "style": {
                "header_styles": {"align": "center"}
            }
        },
        {
            "type": "table",
            "content": {
                "headers": ["Right Test"],
                "rows": [["Data"]]
            },
            "style": {
                "header_styles": {"align": "right"}
            }
        },
        {
            "type": "table",
            "content": {
                "headers": ["Justified Test Header"],
                "rows": [["Data"]]
            },
            "style": {
                "header_styles": {"align": "justify"}
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that tables were created
    tables = doc2.tables
    assert len(tables) == 4
    
    # Check first table header alignment (left)
    header_cell = tables[0].cell(0, 0)
    assert header_cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    
    # Check second table header alignment (center)
    header_cell = tables[1].cell(0, 0)
    assert header_cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    
    # Check third table header alignment (right)
    header_cell = tables[2].cell(0, 0)
    assert header_cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    
    # Check fourth table header alignment (justify)
    header_cell = tables[3].cell(0, 0)
    assert header_cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

def test_table_column_alignment(tmp_path):
    """Test table column alignment for all alignment options including justify"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Col 1", "Col 2", "Col 3", "Col 4"],
                "rows": [
                    ["Left text", "Center text", "Right text", "Justified text for testing"],
                    ["More left", "More center", "More right", "More justified text for proper testing"]
                ]
            },
            "style": {
                "column_styles": {
                    0: {"align": "left"},
                    1: {"align": "center"},
                    2: {"align": "right"},
                    3: {"align": "justify"}
                }
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1
    table = tables[0]
    
    # Check first row alignments
    assert table.cell(1, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert table.cell(1, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert table.cell(1, 2).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert table.cell(1, 3).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Check second row alignments
    assert table.cell(2, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert table.cell(2, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert table.cell(2, 2).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert table.cell(2, 3).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

def test_table_row_alignment(tmp_path):
    """Test table row alignment for all alignment options including justify"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Column 1", "Column 2"],
                "rows": [
                    ["Left row cell 1", "Left row cell 2"],
                    ["Center row cell 1", "Center row cell 2"],
                    ["Right row cell 1", "Right row cell 2"],
                    ["Justified row cell 1", "Justified row cell 2"]
                ]
            },
            "style": {
                "row_styles": {
                    0: {"align": "left"},
                    1: {"align": "center"},
                    2: {"align": "right"},
                    3: {"align": "justify"}
                }
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1
    table = tables[0]
    
    # Check row alignments (row 0 is headers, data rows start at 1)
    # Left row
    assert table.cell(1, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert table.cell(1, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    
    # Center row
    assert table.cell(2, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert table.cell(2, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    
    # Right row
    assert table.cell(3, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert table.cell(3, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    
    # Justified row
    assert table.cell(4, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert table.cell(4, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

def test_table_cell_alignment(tmp_path):
    """Test individual table cell alignment for all alignment options including justify"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Header 1", "Header 2"],
                "rows": [
                    ["Cell (0,0) left", "Cell (0,1) center"],
                    ["Cell (1,0) right", "Cell (1,1) justified"]
                ]
            },
            "style": {
                "cell_styles": {
                    (0, 0): {"align": "left"},
                    (0, 1): {"align": "center"},
                    (1, 0): {"align": "right"},
                    (1, 1): {"align": "justify"}
                }
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1
    table = tables[0]
    
    # Check individual cell alignments (row 0 is headers, data rows start at 1)
    assert table.cell(1, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert table.cell(1, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert table.cell(2, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert table.cell(2, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

def test_text_with_newlines_alignment(tmp_path):
    """Test that alignment works correctly with text blocks containing newlines"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "text",
            "text": "First line\nSecond line\nThird line",
            "style": {"align": "center"}
        },
        {
            "type": "text",
            "text": "Justified paragraph one\n\nJustified paragraph two with longer text to demonstrate justification",
            "style": {"align": "justify"}
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that paragraphs exist and have correct alignment
    paragraphs = [p for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 5  # 3 lines + 2 lines (blank paragraphs are excluded)
    
    # First block - all lines should be center aligned
    assert paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.CENTER
    
    # Second block - all paragraphs should be justified
    assert paragraphs[3].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert paragraphs[4].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

def test_table_cell_with_newlines_alignment(tmp_path):
    """Test that alignment works correctly in table cells containing newlines"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Multi-line"],
                "rows": [
                    ["Line 1\nLine 2\nLine 3"]
                ]
            },
            "style": {
                "column_styles": {
                    0: {"align": "justify"}
                }
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1
    table = tables[0]
    
    # Check that all paragraphs in the cell are justified
    cell = table.cell(1, 0)
    assert len(cell.paragraphs) == 3
    for para in cell.paragraphs:
        assert para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY 