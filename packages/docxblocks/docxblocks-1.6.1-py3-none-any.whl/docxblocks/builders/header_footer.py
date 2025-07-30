"""
Header/Footer Builder Module

This module provides the HeaderFooterBuilder class for setting up headers and footers
in Word documents. It integrates with the existing block system to allow headers and
footers to contain the same content types as the main document.
"""

from docx.enum.section import WD_HEADER_FOOTER
from docxblocks.schema.blocks import HeaderBlock, FooterBlock
from docxblocks.builders.rich_text import RichTextBuilder
import os


class HeaderFooterBuilder:
    """
    Builder class for configuring document headers and footers.
    
    This builder handles the setup of headers and footers across document sections,
    supporting different configurations for first page, odd/even pages, and all pages.
    Header and footer content can include any of the standard block types supported
    by the rich text builder.
    
    Attributes:
        doc: The python-docx Document object
    """
    
    def __init__(self, doc):
        """
        Initialize the HeaderFooterBuilder.
        
        Args:
            doc: The python-docx Document object
        """
        self.doc = doc

    def build_header(self, header_block: HeaderBlock, section_index: int = 0):
        """
        Build a header for the specified section.
        
        Args:
            header_block: A validated HeaderBlock object containing header configuration
            section_index: Index of the section to apply the header to (default: 0)
        """
        section = self.doc.sections[section_index]
        
        # Determine which header(s) to configure based on apply_to setting
        if header_block.apply_to == "all":
            self._configure_header(section.header, header_block.content)
        elif header_block.apply_to == "first":
            section.different_first_page_header_footer = True
            self._configure_header(section.first_page_header, header_block.content)
        elif header_block.apply_to == "all_except_first":
            # Enable different first page but leave first page header empty
            section.different_first_page_header_footer = True
            # Configure header for all pages except first (default header applies to pages 2+)
            self._configure_header(section.header, header_block.content)
        elif header_block.apply_to == "odd":
            # Set up odd/even headers - default header is used for odd pages
            self.doc.settings.odd_and_even_pages_header_footer = True
            self._configure_header(section.header, header_block.content)
        elif header_block.apply_to == "even":
            # Set up odd/even headers - even header is used for even pages
            self.doc.settings.odd_and_even_pages_header_footer = True
            self._configure_header(section.even_page_header, header_block.content)

    def build_footer(self, footer_block: FooterBlock, section_index: int = 0):
        """
        Build a footer for the specified section.
        
        Args:
            footer_block: A validated FooterBlock object containing footer configuration
            section_index: Index of the section to apply the footer to (default: 0)
        """
        section = self.doc.sections[section_index]
        
        # Determine which footer(s) to configure based on apply_to setting
        if footer_block.apply_to == "all":
            self._configure_footer(section.footer, footer_block.content)
        elif footer_block.apply_to == "first":
            section.different_first_page_header_footer = True
            self._configure_footer(section.first_page_footer, footer_block.content)
        elif footer_block.apply_to == "all_except_first":
            # Enable different first page but leave first page footer empty
            section.different_first_page_header_footer = True
            # Configure footer for all pages except first (default footer applies to pages 2+)
            self._configure_footer(section.footer, footer_block.content)
        elif footer_block.apply_to == "odd":
            # Set up odd/even footers - default footer is used for odd pages
            self.doc.settings.odd_and_even_pages_header_footer = True
            self._configure_footer(section.footer, footer_block.content)
        elif footer_block.apply_to == "even":
            # Set up odd/even footers - even footer is used for even pages
            self.doc.settings.odd_and_even_pages_header_footer = True
            self._configure_footer(section.even_page_footer, footer_block.content)

    def _configure_header(self, header, content_blocks):
        """
        Configure a specific header with the provided content blocks.
        
        Args:
            header: The python-docx header object (_Header)
            content_blocks: List of block dictionaries to render in the header
        """
        # Clear any existing content
        for paragraph in header.paragraphs:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        
        # Process content blocks directly using header methods
        for block in content_blocks:
            block_type = block.get('type')
            
            if block_type == 'text':
                para = header.add_paragraph()
                run = para.add_run(block.get('text', ''))
                # Apply text styling if provided
                style = block.get('style', {})
                if style.get('bold'):
                    run.bold = True
                if style.get('italic'):
                    run.italic = True
                if style.get('font_color'):
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor.from_string(style['font_color'])
                if style.get('align'):
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    align_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if style['align'] in align_map:
                        para.alignment = align_map[style['align']]
                        
            elif block_type == 'image':
                para = header.add_paragraph()
                run = para.add_run()
                
                image_path = block.get('path', '')
                style = block.get('style', {})
                
                if image_path and os.path.isfile(image_path):
                    try:
                        from PIL import Image as PILImage
                        from docx.shared import Inches
                        
                        with PILImage.open(image_path) as img:
                            width_px, height_px = img.size
                            dpi_x, dpi_y = img.info.get("dpi", (72, 72))
                            width_in = width_px / dpi_x
                            height_in = height_px / dpi_y

                            # Calculate scale factors for width and height constraints
                            scales = []
                            max_width = self._parse_measurement(style.get("max_width"))
                            max_height = self._parse_measurement(style.get("max_height"))

                            if max_width:
                                scales.append(max_width / width_in)
                            if max_height:
                                scales.append(max_height / height_in)
                            
                            # Use the minimum scale to ensure neither dimension exceeds its maximum
                            if scales:
                                scale = min(scales)
                            else:
                                scale = 1.0

                            # Use file object to ensure embedding works in headers/footers
                            with open(image_path, 'rb') as img_file:
                                run.add_picture(img_file, width=Inches(width_in * scale), height=Inches(height_in * scale))
                                
                    except Exception as e:
                        # Add placeholder text if image fails
                        run.text = "Image could not be loaded"
                else:
                    # Add placeholder text for missing image
                    run.text = "Image not found"
                    
            elif block_type == 'heading':
                para = header.add_paragraph()
                run = para.add_run(block.get('text', ''))
                run.bold = True
                # Apply heading level styling
                level = block.get('level', 1)
                if level == 1:
                    run.font.size = Inches(0.3)  # 18pt
                elif level == 2:
                    run.font.size = Inches(0.25)  # 14pt
                else:
                    run.font.size = Inches(0.2)  # 12pt
                    
            elif block_type == 'table':
                # For tables, we'll need to implement table creation in headers
                # This is more complex and may require a different approach
                para = header.add_paragraph()
                para.add_run("Table in header - not yet implemented")
                
            else:
                # For other block types, add a placeholder
                para = header.add_paragraph()
                para.add_run(f"Block type '{block_type}' not yet supported in headers")

    def _parse_measurement(self, value):
        """
        Parse measurement strings and convert to inches.
        
        Accepts strings like '4in' or '300px' and returns inches as float.
        Supports only inches and pixels for now.
        
        Args:
            value: Measurement string (e.g., "4in", "300px") or None
            
        Returns:
            float: Measurement in inches, or None if parsing fails
        """
        if not value or not isinstance(value, str):
            return None

        value = value.strip().lower()

        if value.endswith("in"):
            try:
                return float(value.replace("in", ""))
            except ValueError:
                return None
        elif value.endswith("px"):
            try:
                px = float(value.replace("px", ""))
                return px / 96.0  # assuming 96 dpi standard
            except ValueError:
                return None
        else:
            return None

    def _configure_footer(self, footer, content_blocks):
        """
        Configure a specific footer with the provided content blocks.
        
        Args:
            footer: The python-docx footer object (_Footer)
            content_blocks: List of block dictionaries to render in the footer
        """
        # Clear any existing content
        for paragraph in footer.paragraphs:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        
        # Process content blocks directly using footer methods
        for block in content_blocks:
            block_type = block.get('type')
            
            if block_type == 'text':
                para = footer.add_paragraph()
                run = para.add_run(block.get('text', ''))
                # Apply text styling if provided
                style = block.get('style', {})
                if style.get('bold'):
                    run.bold = True
                if style.get('italic'):
                    run.italic = True
                if style.get('font_color'):
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor.from_string(style['font_color'])
                if style.get('align'):
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    align_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if style['align'] in align_map:
                        para.alignment = align_map[style['align']]
                        
            elif block_type == 'image':
                para = footer.add_paragraph()
                run = para.add_run()
                
                image_path = block.get('path', '')
                style = block.get('style', {})
                
                if image_path and os.path.isfile(image_path):
                    try:
                        from PIL import Image as PILImage
                        from docx.shared import Inches
                        
                        with PILImage.open(image_path) as img:
                            width_px, height_px = img.size
                            dpi_x, dpi_y = img.info.get("dpi", (72, 72))
                            width_in = width_px / dpi_x
                            height_in = height_px / dpi_y

                            # Calculate scale factors for width and height constraints
                            scales = []
                            max_width = self._parse_measurement(style.get("max_width"))
                            max_height = self._parse_measurement(style.get("max_height"))

                            if max_width:
                                scales.append(max_width / width_in)
                            if max_height:
                                scales.append(max_height / height_in)
                            
                            # Use the minimum scale to ensure neither dimension exceeds its maximum
                            if scales:
                                scale = min(scales)
                            else:
                                scale = 1.0

                            # Use file object to ensure embedding works in headers/footers
                            with open(image_path, 'rb') as img_file:
                                run.add_picture(img_file, width=Inches(width_in * scale), height=Inches(height_in * scale))
                                
                    except Exception as e:
                        # Add placeholder text if image fails
                        run.text = "Image could not be loaded"
                else:
                    # Add placeholder text for missing image
                    run.text = "Image not found"
                    
            elif block_type == 'heading':
                para = footer.add_paragraph()
                run = para.add_run(block.get('text', ''))
                run.bold = True
                # Apply heading level styling
                level = block.get('level', 1)
                if level == 1:
                    run.font.size = Inches(0.3)  # 18pt
                elif level == 2:
                    run.font.size = Inches(0.25)  # 14pt
                else:
                    run.font.size = Inches(0.2)  # 12pt
                    
            elif block_type == 'table':
                # For tables, we'll need to implement table creation in footers
                # This is more complex and may require a different approach
                para = footer.add_paragraph()
                para.add_run("Table in footer - not yet implemented")
                
            else:
                # For other block types, add a placeholder
                para = footer.add_paragraph()
                para.add_run(f"Block type '{block_type}' not yet supported in footers") 