import os
from docx import Document
from docxblocks.core.inserter import DocxBuilder

def test_inline_text_default(tmp_path):
    """Test that text blocks each create their own paragraph (no grouping)"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "Participant Name: "},
        {"type": "text", "text": "John Doe"},
        {"type": "text", "text": " (ID: "},
        {"type": "text", "text": "12345"},
        {"type": "text", "text": ")"},
        {"type": "text", "text": "New line starts here", "new_paragraph": True},
        {"type": "text", "text": "This should be on the same line as above"},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Each block is now its own paragraph
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 7
    assert "Participant Name: " in paragraphs[0]
    assert "John Doe" in paragraphs[1]
    assert " (ID: " in paragraphs[2]
    assert "12345" in paragraphs[3]
    assert ")" in paragraphs[4]
    assert "New line starts here" in paragraphs[5]
    assert "This should be on the same line as above" in paragraphs[6]

def test_mixed_inline_and_paragraphs(tmp_path):
    """Test mixing inline text with other block types, all text blocks are separate paragraphs"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "Status: "},
        {"type": "text", "text": "Active"},
        {"type": "heading", "text": "Summary", "level": 2},
        {"type": "text", "text": "This is a "},
        {"type": "text", "text": "summary", "style": {"bold": True}},
        {"type": "text", "text": " of the report."},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Each text block is now its own paragraph
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 6
    assert "Status: " in paragraphs[0]
    assert "Active" in paragraphs[1]
    assert "Summary" in paragraphs[2]
    assert "This is a " in paragraphs[3]
    assert "summary" in paragraphs[4]
    assert " of the report." in paragraphs[5]

def test_inline_after_new_paragraph(tmp_path):
    """Test that each text block after a new_paragraph is its own paragraph"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "Participant No: ", "style": {"bold": True}, "new_paragraph": True},
        {"type": "text", "text": "12345", "style": {"bold": True}},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Each block is now its own paragraph
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 2
    assert "Participant No: " in paragraphs[0]
    assert "12345" in paragraphs[1] 