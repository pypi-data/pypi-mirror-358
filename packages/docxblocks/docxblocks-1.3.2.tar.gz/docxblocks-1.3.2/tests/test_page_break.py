import os
from docx import Document
from docx.enum.text import WD_BREAK
from docxblocks.core.inserter import DocxBuilder

def test_page_break_simple(tmp_path):
    """Simple test that page breaks are created without errors"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "Before page break"},
        {"type": "page_break"},
        {"type": "text", "text": "After page break"},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Should have 3 paragraphs total
    all_paragraphs = list(doc2.paragraphs)
    assert len(all_paragraphs) == 3
    
    # Should have 2 text paragraphs
    text_paragraphs = [p for p in doc2.paragraphs if p.text.strip()]
    assert len(text_paragraphs) == 2
    assert "Before page break" in text_paragraphs[0].text
    assert "After page break" in text_paragraphs[1].text

def test_page_break_with_inline_text(tmp_path):
    """Test that page breaks work correctly with each text block as its own paragraph"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "First page: "},
        {"type": "text", "text": "inline text"},
        {"type": "page_break"},
        {"type": "text", "text": "Second page: "},
        {"type": "text", "text": "more inline text"},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Each text block is now its own paragraph
    all_paragraphs = list(doc2.paragraphs)
    assert len(all_paragraphs) == 5  # Four text paragraphs + page break
    assert "First page: " in all_paragraphs[0].text
    assert "inline text" in all_paragraphs[1].text
    assert "Second page: " in all_paragraphs[3].text
    assert "more inline text" in all_paragraphs[4].text 