import os
from docx import Document
from docxblocks.core.inserter import DocxBuilder
from docxblocks.constants import DEFAULT_EMPTY_VALUE_TEXT

def test_text_block(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "", "style": {"align": "center"}}
    ]
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    found = any(DEFAULT_EMPTY_VALUE_TEXT in p.text for p in doc2.paragraphs)
    assert found, "Placeholder text not found in output docx"

def test_single_newline_creates_paragraph(tmp_path):
    """Test that single newlines create new paragraphs"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "First line\nSecond line"},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
    
    # Should have 2 paragraphs from the single newline
    assert len(paragraphs) == 2
    assert "First line" in paragraphs[0]
    assert "Second line" in paragraphs[1]

def test_double_newline_creates_paragraph_with_blank_line(tmp_path):
    """Test that double newlines create two separate paragraphs"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "First paragraph\n\nSecond paragraph"},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
    
    # Should have 3 paragraphs: First, empty, Second
    assert len(paragraphs) == 2
    assert "First paragraph" in paragraphs[0]
    assert "Second paragraph" in paragraphs[1]
    
    # Count total paragraphs (including empty ones)
    all_paragraphs = [p.text for p in doc2.paragraphs]
    
    # Should have 3 total paragraphs: First, empty, Second
    assert len(all_paragraphs) == 3
    assert all_paragraphs[1] == ""  # Empty paragraph from \n\n

def test_mixed_newlines(tmp_path):
    """Test mixed single and double newlines"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "First\nSecond\n\nThird\nFourth"},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
    
    # Should have 4 paragraphs with text: First, Second, Third, Fourth
    assert len(paragraphs) == 4
    assert "First" in paragraphs[0]
    assert "Second" in paragraphs[1]
    assert "Third" in paragraphs[2]
    assert "Fourth" in paragraphs[3]
    
    # Count total paragraphs (including empty ones)
    all_paragraphs = [p.text for p in doc2.paragraphs]
    
    # Should have 5 total paragraphs: First, Second, empty, Third, Fourth
    assert len(all_paragraphs) == 5
    assert all_paragraphs[2] == ""  # Empty paragraph from \n\n

def test_spacing_functionality(tmp_path):
    """Test that spacing parameter adds extra blank lines after new_paragraph text"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "First paragraph", "new_paragraph": True},
        {"type": "text", "text": "Second paragraph", "new_paragraph": True, "spacing": 2},
        {"type": "text", "text": "Third paragraph", "new_paragraph": True},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
    
    # Should have 3 paragraphs with text
    assert len(paragraphs) == 3
    assert "First paragraph" in paragraphs[0]
    assert "Second paragraph" in paragraphs[1]
    assert "Third paragraph" in paragraphs[2]
    
    # Count total paragraphs (including blank ones)
    all_paragraphs = [p.text for p in doc2.paragraphs]
    
    # Should have 4 total paragraphs: First, Second, blank line, Third
    assert len(all_paragraphs) == 4
    assert all_paragraphs[2] == ""  # Blank line from spacing: 2

def test_newlines_create_paragraphs(tmp_path):
    """Test that all newlines create new paragraphs"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "Line 1\nLine 2\nLine 3"},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
    
    # Should have 3 paragraphs from the single text block with newlines
    assert len(paragraphs) == 3
    assert "Line 1" in paragraphs[0]
    assert "Line 2" in paragraphs[1]
    assert "Line 3" in paragraphs[2] 