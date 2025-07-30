"""
Test inline text functionality
"""

import os
from docx import Document
from docxblocks import DocxBuilder


def test_inline_text_default(tmp_path):
    """Test that text blocks are inline by default (no new paragraphs)"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "Participant Name: "},
        {"type": "text", "text": "John Doe"},
        {"type": "text", "text": " (ID: "},
        {"type": "text", "text": "12345"},
        {"type": "text", "text": ")"},
        {"type": "text", "text": "\nNew line starts here"},
        {"type": "text", "text": "This should be on the same line as above"},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Should have 2 paragraphs: inline text, new paragraph text
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 2
    assert "Participant Name: John Doe (ID: 12345)" in paragraphs[0]
    assert "New line starts hereThis should be on the same line as above" in paragraphs[1]


def test_mixed_inline_and_paragraphs(tmp_path):
    """Test mixing inline text with paragraph breaks"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "First paragraph "},
        {"type": "text", "text": "continues inline"},
        {"type": "text", "text": "\nSecond paragraph "},
        {"type": "text", "text": "also continues inline"},
        {"type": "text", "text": "\n\nThird paragraph with blank line above"},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Should have 3 paragraphs with text
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 3
    assert "First paragraph continues inline" in paragraphs[0]
    assert "Second paragraph also continues inline" in paragraphs[1]
    assert "Third paragraph with blank line above" in paragraphs[2]


def test_inline_after_new_paragraph(tmp_path):
    """Test that inline text works correctly after a new paragraph block"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "\nParticipant No: ", "style": {"bold": True}},
        {"type": "text", "text": "12345", "style": {"bold": True}},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Should have 1 paragraph with text (grouped inline)
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 1
    assert "Participant No: 12345" in paragraphs[0] 