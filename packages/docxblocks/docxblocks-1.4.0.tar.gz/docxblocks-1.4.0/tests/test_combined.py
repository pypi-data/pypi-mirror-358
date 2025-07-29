import os
from docx import Document
from docxblocks.core.inserter import DocxBuilder

def test_combined(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "heading", "text": "Title", "level": 1},
        {"type": "text", "text": "Summary."},
        {"type": "bullets", "items": ["A", "B", "C"]},
        {"type": "table", "content": {"headers": ["H1", "H2"], "rows": [["1", "2"]]}},
        {"type": "image", "path": "nonexistent.png"}
    ]
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    # Check for heading, text, bullets, table, and image placeholder
    texts = [p.text for p in doc2.paragraphs]
    assert any("Title" in t for t in texts)
    assert any("Summary" in t for t in texts)
    # Check for bullet items' text
    for bullet_item in ["A", "B", "C"]:
        assert any(bullet_item == t.strip() for t in texts)
    # Check for bullet paragraph style (Word-native bullets)
    bullet_styles = [p.style.name for p in doc2.paragraphs if p.text.strip() in ["A", "B", "C"] and p.style is not None]
    assert any(s in ("List Bullet", "DocxBlocks_Bullet") for s in bullet_styles)
    assert len(doc2.tables) == 1 