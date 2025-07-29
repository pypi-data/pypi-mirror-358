"""
Test block order preservation
"""

import os
from docx import Document
from docxblocks import DocxBuilder


def test_block_order_preservation(tmp_path):
    """Test that blocks are inserted in the correct order"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "heading", "text": "First Heading", "level": 1},
        {"type": "text", "text": "First paragraph"},
        {"type": "text", "text": "\nSecond paragraph"},
        {"type": "heading", "text": "Second Heading", "level": 2},
        {"type": "text", "text": "Third paragraph"},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Check that blocks are in the correct order
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
    assert "First Heading" in paragraphs[0]
    assert "First paragraph" in paragraphs[1]
    assert "Second paragraph" in paragraphs[2]
    assert "Second Heading" in paragraphs[3]
    assert "Third paragraph" in paragraphs[4]


def test_mixed_block_types_order(tmp_path):
    """Test that mixed block types maintain correct order"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "heading", "text": "Report", "level": 1},
        {"type": "text", "text": "Introduction"},
        {"type": "table", "content": {
            "headers": ["Name", "Value"],
            "rows": [["Item 1", "100"], ["Item 2", "200"]]
        }},
        {"type": "text", "text": "Conclusion"},
        {"type": "bullets", "items": ["Point 1", "Point 2"]},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Check that blocks are in the correct order
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
    assert "Report" in paragraphs[0]
    assert "Introduction" in paragraphs[1]
    assert "Conclusion" in paragraphs[2]
    assert "Point 1" in paragraphs[3]
    assert "Point 2" in paragraphs[4]


def test_inline_text_order(tmp_path):
    """Test that inline text blocks maintain their order and grouping"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    # Create inline text blocks that should stay in order
    blocks = [
        {"type": "text", "text": "First "},
        {"type": "text", "text": "Second "},
        {"type": "text", "text": "\nThird"},
        {"type": "text", "text": "Fourth "},
        {"type": "text", "text": "\nFifth"},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
    # Should have: "First Second", "ThirdFourth", "Fifth"
    assert "First Second" in paragraphs[0]
    assert "ThirdFourth" in paragraphs[1]
    assert "Fifth" in paragraphs[2] 