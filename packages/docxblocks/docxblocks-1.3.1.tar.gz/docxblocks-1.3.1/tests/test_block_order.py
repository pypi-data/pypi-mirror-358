import os
from docx import Document
from docxblocks.core.inserter import DocxBuilder

def test_block_order_preservation(tmp_path):
    """Test that blocks are processed in the exact order they are provided"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    # Create blocks in a specific order
    blocks = [
        {"type": "heading", "text": "First Block", "level": 1},
        {"type": "text", "text": "Second Block", "new_paragraph": True},
        {"type": "heading", "text": "Third Block", "level": 2},
        {"type": "text", "text": "Fourth Block", "new_paragraph": True},
        {"type": "bullets", "items": ["Fifth Block - Item 1", "Fifth Block - Item 2"]},
        {"type": "text", "text": "Sixth Block", "new_paragraph": True},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
    
    # Verify the order matches our expected sequence
    expected_order = [
        "First Block",  # Heading 1
        "Second Block",  # Text
        "Third Block",   # Heading 2
        "Fourth Block",  # Text
        "Fifth Block - Item 1",  # First bullet
        "Fifth Block - Item 2",  # Second bullet
        "Sixth Block",   # Text
    ]
    
    # Check that we have the right number of paragraphs
    assert len(paragraphs) == len(expected_order), f"Expected {len(expected_order)} paragraphs, got {len(paragraphs)}"
    
    # Check that each paragraph appears in the correct order
    for i, expected in enumerate(expected_order):
        assert expected in paragraphs[i], f"Expected '{expected}' at position {i}, but found '{paragraphs[i]}'"

def test_mixed_block_types_order(tmp_path):
    """Test that mixed block types maintain their order"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    # Create blocks with mixed types in specific order
    blocks = [
        {"type": "text", "text": "Start", "new_paragraph": True},
        {"type": "heading", "text": "Section 1", "level": 1},
        {"type": "text", "text": "Content 1", "new_paragraph": True},
        {"type": "table", "content": {"headers": ["Col1"], "rows": [["Data1"]]}},
        {"type": "text", "text": "Content 2", "new_paragraph": True},
        {"type": "heading", "text": "Section 2", "level": 2},
        {"type": "bullets", "items": ["Item 1", "Item 2"]},
        {"type": "text", "text": "End", "new_paragraph": True},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
    
    # Verify the order - we should see text and headings in order
    expected_text_order = [
        "Start",
        "Section 1", 
        "Content 1",
        "Content 2",
        "Section 2",
        "Item 1",
        "Item 2", 
        "End"
    ]
    
    # Check that we have the right number of text paragraphs
    assert len(paragraphs) == len(expected_text_order), f"Expected {len(expected_text_order)} text paragraphs, got {len(paragraphs)}"
    
    # Check that each text paragraph appears in the correct order
    for i, expected in enumerate(expected_text_order):
        assert expected in paragraphs[i], f"Expected '{expected}' at position {i}, but found '{paragraphs[i]}'"
    
    # Also verify that tables are present (they don't add text paragraphs)
    tables = doc2.tables
    assert len(tables) == 1, "Expected 1 table to be present"

def test_inline_text_order(tmp_path):
    """Test that inline text blocks maintain their order and grouping"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    # Create inline text blocks that should stay in order
    blocks = [
        {"type": "text", "text": "First "},
        {"type": "text", "text": "Second "},
        {"type": "text", "text": "Third", "new_paragraph": True},
        {"type": "text", "text": "Fourth "},
        {"type": "text", "text": "Fifth", "new_paragraph": True},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]

    # The text builder creates separate paragraphs for each new_paragraph=True block
    # So we should have: "First Second", "Third", "Fourth", "Fifth"
    assert len(paragraphs) == 4
    assert "First Second" in paragraphs[0]
    assert "Third" in paragraphs[1]
    assert "Fourth" in paragraphs[2]
    assert "Fifth" in paragraphs[3] 