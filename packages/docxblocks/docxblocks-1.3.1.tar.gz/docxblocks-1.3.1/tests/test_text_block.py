import os
from docx import Document
from docxblocks.core.inserter import DocxBuilder
from docxblocks.constants import DEFAULT_EMPTY_VALUE_TEXT

def test_text_block(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "", "style": {"align": "center"}}
    ]
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    found = any(DEFAULT_EMPTY_VALUE_TEXT in p.text for p in doc2.paragraphs)
    assert found, "Placeholder text not found in output docx"

def test_double_newline_creates_paragraph_with_blank_line(tmp_path):
    """Test that \n\n creates a new paragraph with a blank line before it"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "First paragraph.\n\nSecond paragraph."},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Should have 3 paragraphs: first paragraph, blank line, second paragraph
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 2
    
    # First paragraph should contain the first part
    assert "First paragraph." in paragraphs[0]
    
    # Second paragraph should contain the second part
    assert "Second paragraph." in paragraphs[1]
    
    # Check that there's a blank paragraph between them
    all_paragraphs = list(doc2.paragraphs)
    assert len(all_paragraphs) == 3  # First, blank, second
    
    # The middle paragraph should be blank
    assert not all_paragraphs[1].text.strip()

def test_single_newline_remains_inline(tmp_path):
    """Test that single \n remains as inline text"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "Line 1\nLine 2"},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Should have only 1 paragraph with both lines
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 1
    
    # The paragraph should contain both lines with literal \n
    assert "Line 1\nLine 2" in paragraphs[0]

def test_mixed_newlines(tmp_path):
    """Test mixing single and double newlines"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "First line\nSecond line\n\nThird paragraph."},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Should have 2 paragraphs: first two lines together, then third paragraph
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 2
    
    # First paragraph should contain the first two lines
    assert "First line\nSecond line" in paragraphs[0]
    
    # Second paragraph should contain the third line
    assert "Third paragraph." in paragraphs[1]
    
    # Check that there's a blank paragraph between them
    all_paragraphs = list(doc2.paragraphs)
    assert len(all_paragraphs) == 3  # First, blank, second
    
    # The middle paragraph should be blank
    assert not all_paragraphs[1].text.strip() 