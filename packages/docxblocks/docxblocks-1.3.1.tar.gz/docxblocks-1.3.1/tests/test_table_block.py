import os
from docx import Document
from docxblocks.core.inserter import DocxBuilder

def test_table_block(tmp_path):
    """Test basic table block functionality"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Name", "Status", "Notes"],
                "rows": [
                    ["Service A", "OK", "All systems operational"],
                    ["Service B", "Warning", "High memory usage"],
                    ["Service C", "Error", "Connection timeout"]
                ]
            },
            "style": {
                "header_styles": {"bold": True, "bg_color": "f2f2f2"},
                "column_widths": [0.3, 0.2, 0.5]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1
    
    table = tables[0]
    assert len(table.rows) == 4  # 1 header + 3 data rows
    assert len(table.columns) == 3

def test_table_block_with_integers(tmp_path):
    """Test table block with integer values"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["ID", "Count", "Percentage"],
                "rows": [
                    [1, 100, 25],
                    [2, 200, 50],
                    [3, 100, 25]
                ]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created with integer values
    tables = doc2.tables
    assert len(tables) == 1
    
    table = tables[0]
    # Check that integer values are properly converted to strings
    assert table.cell(1, 0).text == "1"
    assert table.cell(1, 1).text == "100"
    assert table.cell(1, 2).text == "25"

def test_table_cell_with_single_newlines(tmp_path):
    """Test that single \n in table cells remains as literal newlines"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Name", "Description"],
                "rows": [
                    ["Item 1", "Line 1\nLine 2\nLine 3"],
                    ["Item 2", "Single line"]
                ]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1
    
    table = tables[0]
    
    # Check first row - should have single paragraph with literal newlines
    first_cell = table.cell(1, 1)  # Row 1, Column 1 (description)
    paragraphs = first_cell.paragraphs
    
    # Should have only 1 paragraph with literal newlines
    assert len(paragraphs) == 1
    
    # The paragraph should contain all lines with literal \n
    assert "Line 1\nLine 2\nLine 3" in paragraphs[0].text

def test_table_cell_with_double_newlines(tmp_path):
    """Test that \n\n in table cells creates new paragraphs with blank lines"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Name", "Description"],
                "rows": [
                    ["Item 1", "First paragraph.\n\nSecond paragraph with blank line."],
                    ["Item 2", "Single line\nAnother line\n\nNew paragraph."]
                ]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1
    
    table = tables[0]
    
    # Check first row - should have multiple paragraphs in the description cell
    first_cell = table.cell(1, 1)  # Row 1, Column 1 (description)
    paragraphs = first_cell.paragraphs
    
    # Should have 3 paragraphs: first paragraph, blank line, second paragraph
    assert len(paragraphs) == 3
    
    # First paragraph should contain the first part
    assert "First paragraph." in paragraphs[0].text
    
    # Second paragraph should be blank
    assert not paragraphs[1].text.strip()
    
    # Third paragraph should contain the second part
    assert "Second paragraph with blank line." in paragraphs[2].text
    
    # Check second row - should have mixed single and double newlines
    second_cell = table.cell(2, 1)  # Row 2, Column 1 (description)
    paragraphs = second_cell.paragraphs
    
    # Should have 3 paragraphs: first two lines together, blank line, new paragraph
    assert len(paragraphs) == 3
    
    # First paragraph should contain the first two lines
    assert "Single line\nAnother line" in paragraphs[0].text
    
    # Second paragraph should be blank
    assert not paragraphs[1].text.strip()
    
    # Third paragraph should contain the new paragraph
    assert "New paragraph." in paragraphs[2].text

def test_table_header_with_double_newlines(tmp_path):
    """Test that \n\n in table headers creates new paragraphs with blank lines"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Name", "Description\n\nDetails"],
                "rows": [
                    ["Item 1", "Some description"],
                    ["Item 2", "Another description"]
                ]
            }
        }
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Check that table was created
    tables = doc2.tables
    assert len(tables) == 1
    
    table = tables[0]
    
    # Check header cell with double newlines
    header_cell = table.cell(0, 1)  # Header row, Column 1
    paragraphs = header_cell.paragraphs
    
    # Should have 3 paragraphs: first part, blank line, second part
    assert len(paragraphs) == 3
    
    # First paragraph should contain the first part
    assert "Description" in paragraphs[0].text
    
    # Second paragraph should be blank
    assert not paragraphs[1].text.strip()
    
    # Third paragraph should contain the second part
    assert "Details" in paragraphs[2].text 