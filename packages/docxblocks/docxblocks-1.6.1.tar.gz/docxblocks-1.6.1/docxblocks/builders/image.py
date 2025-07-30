"""
Image Builder Module

This module provides the ImageBuilder class for rendering image blocks in Word documents.
It handles image insertion with automatic sizing, DPI calculation, and error handling.
"""

from docx.shared import Inches, RGBColor
from PIL import Image as PILImage
import os
from docxblocks.constants import DEFAULT_EMPTY_VALUE_TEXT, DEFAULT_EMPTY_VALUE_STYLE

class ImageBuilder:
    """
    Builder class for rendering image blocks in Word documents.
    
    This builder handles image insertion with automatic sizing based on DPI,
    support for maximum width/height constraints, and graceful error handling
    for missing or invalid image files.
    
    The builder automatically calculates image dimensions from DPI information
    and supports scaling based on maximum width/height constraints. Invalid
    or missing images are replaced with consistent placeholders.
    """
    
    @staticmethod
    def build(doc, image_path=None, parent=None, index=None, **style_kwargs):
        """
        Build and render an image block in the document.
        
        This method processes the image path, validates file existence,
        calculates appropriate dimensions, and inserts the image into the
        document. Handles errors gracefully with placeholder text.
        
        Args:
            doc: The python-docx Document object
            image_path: Path to the image file (can be None, empty, or invalid)
            parent: The parent XML element where content will be inserted
            index: The insertion index within the parent element
            **style_kwargs: Additional styling options including:
                - max_width: Maximum width constraint (e.g., "4in", "300px")
                - max_height: Maximum height constraint (e.g., "4in", "300px")
        """
        # Validate required parameters
        if parent is None or index is None:
            return
            
        # Handle empty or invalid image path with placeholder
        if not image_path or not image_path.strip() or not os.path.isfile(image_path):
            if hasattr(parent, 'add_paragraph'):
                para = parent.add_paragraph(DEFAULT_EMPTY_VALUE_TEXT)
            else:
                para = doc.add_paragraph(DEFAULT_EMPTY_VALUE_TEXT)
            # Apply placeholder style
            run = para.runs[0]
            run.font.bold = DEFAULT_EMPTY_VALUE_STYLE.get("bold", True)
            if DEFAULT_EMPTY_VALUE_STYLE.get("font_color"):
                run.font.color.rgb = RGBColor.from_string(DEFAULT_EMPTY_VALUE_STYLE["font_color"])
            parent.insert(index, para._element)
            return

        if hasattr(parent, 'add_paragraph'):
            new_para = parent.add_paragraph()
        else:
            new_para = doc.add_paragraph()
        run = new_para.add_run()

        try:
            with PILImage.open(image_path) as img:
                width_px, height_px = img.size
                dpi_x, dpi_y = img.info.get("dpi", (72, 72))
                width_in = width_px / dpi_x
                height_in = height_px / dpi_y

                # Calculate scale factors for width and height constraints
                scales = []
                max_width = _parse_measurement(style_kwargs.get("max_width"))
                max_height = _parse_measurement(style_kwargs.get("max_height"))

                if max_width:
                    scales.append(max_width / width_in)
                if max_height:
                    scales.append(max_height / height_in)
                
                # Use the minimum scale to ensure neither dimension exceeds its maximum
                # This allows upscaling if the image is smaller than the specified constraints
                if scales:
                    scale = min(scales)
                else:
                    scale = 1.0

                # Use file object to ensure embedding works in headers/footers
                with open(image_path, 'rb') as img_file:
                    run.add_picture(img_file, width=Inches(width_in * scale), height=Inches(height_in * scale))

        except Exception as e:
            if hasattr(parent, 'add_paragraph'):
                error_para = parent.add_paragraph(DEFAULT_EMPTY_VALUE_TEXT)
            else:
                error_para = doc.add_paragraph(DEFAULT_EMPTY_VALUE_TEXT)
            # Apply placeholder style
            run = error_para.runs[0]
            run.font.bold = DEFAULT_EMPTY_VALUE_STYLE.get("bold", True)
            if DEFAULT_EMPTY_VALUE_STYLE.get("font_color"):
                run.font.color.rgb = RGBColor.from_string(DEFAULT_EMPTY_VALUE_STYLE["font_color"])
            parent.insert(index, error_para._element)
            return

        parent.insert(index, new_para._element)


def _parse_measurement(value):
    """
    Parse measurement strings and convert to inches.
    
    Accepts strings like '4in' or '300px' and returns inches as float.
    Supports only inches and pixels for now.
    
    Args:
        value: Measurement string (e.g., "4in", "300px") or None
        
    Returns:
        float: Measurement in inches, or None if parsing fails
    """
    if not value or not isinstance(value, str):
        return None

    value = value.strip().lower()

    if value.endswith("in"):
        try:
            return float(value.replace("in", ""))
        except ValueError:
            return None
    elif value.endswith("px"):
        try:
            px = float(value.replace("px", ""))
            return px / 96.0  # assuming 96 dpi standard
        except ValueError:
            return None
    else:
        return None
