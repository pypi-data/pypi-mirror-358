import os
from docx import Document
from docxblocks.core.inserter import DocxBuilder

def test_inline_text_default(tmp_path):
    """Test that text blocks are inline by default (no new paragraphs)"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "Participant Name: "},
        {"type": "text", "text": "John Doe"},
        {"type": "text", "text": " (ID: "},
        {"type": "text", "text": "12345"},
        {"type": "text", "text": ")"},
        {"type": "text", "text": "New line starts here", "new_paragraph": True},
        {"type": "text", "text": "This should be on the same line as above"},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Should have 2 paragraphs: inline text, new paragraph text + inline text
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 2
    
    # First paragraph should contain all the inline text
    first_para = paragraphs[0]
    assert "Participant Name: John Doe (ID: 12345)" in first_para
    
    # Second paragraph should contain the new paragraph text and the inline text after it
    second_para = paragraphs[1]
    assert "New line starts hereThis should be on the same line as above" in second_para

def test_mixed_inline_and_paragraphs(tmp_path):
    """Test mixing inline text with other block types"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "Status: "},
        {"type": "text", "text": "Active"},
        {"type": "heading", "text": "Summary", "level": 2},
        {"type": "text", "text": "This is a "},
        {"type": "text", "text": "summary", "style": {"bold": True}},
        {"type": "text", "text": " of the report."},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Should have 3 paragraphs: inline text, heading, inline text
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 3
    
    # First paragraph should contain inline status text
    assert "Status: Active" in paragraphs[0]
    
    # Second paragraph should be the heading
    assert "Summary" in paragraphs[1]
    
    # Third paragraph should contain inline summary text
    assert "This is a summary of the report." in paragraphs[2]

def test_inline_after_new_paragraph(tmp_path):
    """Test that inline text works correctly after a new_paragraph block"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "Participant No: ", "style": {"bold": True}, "new_paragraph": True},
        {"type": "text", "text": "12345", "style": {"bold": True}},
    ]
    
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    
    # Should have only 1 paragraph with both pieces of text
    paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert len(paragraphs) == 1
    
    # The paragraph should contain both pieces of text
    assert "Participant No: 12345" in paragraphs[0] 