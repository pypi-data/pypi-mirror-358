"""
Header and Footer Builder Module

This module provides the HeaderFooterBuilder class for rendering header and footer blocks in Word documents.
It handles header and footer insertion with support for various page application rules and rich content.
"""

from docx import Document
from docxblocks.schema.blocks import HeaderBlock, FooterBlock
from docxblocks.schema.shared import TextStyle
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class HeaderFooterBuilder:
    """
    Builder class for rendering header and footer blocks in Word documents.
    
    This builder handles header and footer insertion with support for various
    page application rules (all pages, all except first, first only, odd/even).
    It supports rich content including text, images, and other block types.
    
    The builder automatically configures section properties to apply headers
    and footers according to the specified rules and handles content rendering
    using the same logic as the main document builder.
    """
    
    def __init__(self, doc):
        """
        Initialize the HeaderFooterBuilder.
        
        Args:
            doc: The python-docx Document object
        """
        self.doc = doc
    
    def build_header(self, block: HeaderBlock):
        """
        Build and render a header block.
        
        Args:
            block: A validated HeaderBlock object
        """
        # Get the section
        section = self.doc.sections[0]
        
        # Configure header based on apply_to rule
        self._configure_header_application(section, block.apply_to)
        
        # Get the appropriate header
        header = self._get_header_for_rule(section, block.apply_to)
        
        # Render content in the header
        self._configure_header(header, block.content)
    
    def build_footer(self, block: FooterBlock):
        """
        Build and render a footer block.
        
        Args:
            block: A validated FooterBlock object
        """
        # Get the section
        section = self.doc.sections[0]
        
        # Configure footer based on apply_to rule
        self._configure_footer_application(section, block.apply_to)
        
        # Get the appropriate footer
        footer = self._get_footer_for_rule(section, block.apply_to)
        
        # Render content in the footer
        self._configure_footer(footer, block.content)
    
    def _configure_header_application(self, section, apply_to):
        """
        Configure header application based on the apply_to rule.
        
        Args:
            section: The document section
            apply_to: The application rule
        """
        if apply_to in ["first", "all_except_first"]:
            section.different_first_page_header_footer = True
        elif apply_to in ["odd", "even"]:
            self.doc.settings.odd_and_even_pages_header_footer = True
    
    def _configure_footer_application(self, section, apply_to):
        """
        Configure footer application based on the apply_to rule.
        
        Args:
            section: The document section
            apply_to: The application rule
        """
        if apply_to in ["first", "all_except_first"]:
            section.different_first_page_header_footer = True
        elif apply_to in ["odd", "even"]:
            self.doc.settings.odd_and_even_pages_header_footer = True
    
    def _get_header_for_rule(self, section, apply_to):
        """
        Get the appropriate header based on the apply_to rule.
        
        Args:
            section: The document section
            apply_to: The application rule
            
        Returns:
            The appropriate header object
        """
        if apply_to == "first":
            return section.first_page_header
        elif apply_to == "odd":
            return section.header
        elif apply_to == "even":
            return section.even_page_header
        else:  # "all" or "all_except_first"
            return section.header
    
    def _get_footer_for_rule(self, section, apply_to):
        """
        Get the appropriate footer based on the apply_to rule.
        
        Args:
            section: The document section
            apply_to: The application rule
            
        Returns:
            The appropriate footer object
        """
        if apply_to == "first":
            return section.first_page_footer
        elif apply_to == "odd":
            return section.footer
        elif apply_to == "even":
            return section.even_page_footer
        else:  # "all" or "all_except_first"
            return section.footer

    def _configure_header(self, header, content_blocks):
        """
        Configure a specific header with the provided content blocks.
        
        Args:
            header: The python-docx header object (_Header)
            content_blocks: List of block dictionaries to render in the header
        """
        # Clear any existing content
        for paragraph in header.paragraphs:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        
        # Process content blocks directly using header methods
        for block in content_blocks:
            block_type = block.get('type')
            
            if block_type == 'text':
                para = header.add_paragraph()
                run = para.add_run(block.get('text', ''))
                # Apply text styling if provided
                style = block.get('style', {})
                if style.get('bold'):
                    run.bold = True
                if style.get('italic'):
                    run.italic = True
                if style.get('font_color'):
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor.from_string(style['font_color'])
                if style.get('align'):
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    align_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if style['align'] in align_map:
                        para.alignment = align_map[style['align']]
                        
            elif block_type == 'image':
                para = header.add_paragraph()
                run = para.add_run()
                
                image_path = block.get('path', '')
                style = block.get('style', {})
                
                if image_path and os.path.isfile(image_path):
                    try:
                        from PIL import Image as PILImage
                        from docx.shared import Inches
                        
                        with PILImage.open(image_path) as img:
                            width_px, height_px = img.size
                            dpi_x, dpi_y = img.info.get("dpi", (72, 72))
                            width_in = width_px / dpi_x
                            height_in = height_px / dpi_y

                            # Calculate scale factors for width and height constraints
                            scales = []
                            max_width = self._parse_measurement(style.get("max_width"))
                            max_height = self._parse_measurement(style.get("max_height"))

                            if max_width:
                                scales.append(max_width / width_in)
                            if max_height:
                                scales.append(max_height / height_in)
                            
                            # Use the minimum scale to ensure neither dimension exceeds its maximum
                            if scales:
                                scale = min(scales)
                            else:
                                scale = 1.0

                            # Use file object to ensure embedding works in headers/footers
                            with open(image_path, 'rb') as img_file:
                                run.add_picture(img_file, width=Inches(width_in * scale), height=Inches(height_in * scale))
                            
                            # Apply text wrapping and positioning properties
                            self._apply_image_wrapping(run, style)
                            
                    except Exception as e:
                        # Handle image loading errors gracefully
                        run.text = f"[Image: {os.path.basename(image_path)}]"
                else:
                    # Handle missing image
                    run.text = f"[Missing Image: {image_path}]"

    def _configure_footer(self, footer, content_blocks):
        """
        Configure a specific footer with the provided content blocks.
        
        Args:
            footer: The python-docx footer object (_Footer)
            content_blocks: List of block dictionaries to render in the footer
        """
        # Clear any existing content
        for paragraph in footer.paragraphs:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        
        # Process content blocks directly using footer methods
        for block in content_blocks:
            block_type = block.get('type')
            
            if block_type == 'text':
                para = footer.add_paragraph()
                run = para.add_run(block.get('text', ''))
                # Apply text styling if provided
                style = block.get('style', {})
                if style.get('bold'):
                    run.bold = True
                if style.get('italic'):
                    run.italic = True
                if style.get('font_color'):
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor.from_string(style['font_color'])
                if style.get('align'):
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    align_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if style['align'] in align_map:
                        para.alignment = align_map[style['align']]
                        
            elif block_type == 'image':
                para = footer.add_paragraph()
                run = para.add_run()
                
                image_path = block.get('path', '')
                style = block.get('style', {})
                
                if image_path and os.path.isfile(image_path):
                    try:
                        from PIL import Image as PILImage
                        from docx.shared import Inches
                        
                        with PILImage.open(image_path) as img:
                            width_px, height_px = img.size
                            dpi_x, dpi_y = img.info.get("dpi", (72, 72))
                            width_in = width_px / dpi_x
                            height_in = height_px / dpi_y

                            # Calculate scale factors for width and height constraints
                            scales = []
                            max_width = self._parse_measurement(style.get("max_width"))
                            max_height = self._parse_measurement(style.get("max_height"))

                            if max_width:
                                scales.append(max_width / width_in)
                            if max_height:
                                scales.append(max_height / height_in)
                            
                            # Use the minimum scale to ensure neither dimension exceeds its maximum
                            if scales:
                                scale = min(scales)
                            else:
                                scale = 1.0

                            # Use file object to ensure embedding works in headers/footers
                            with open(image_path, 'rb') as img_file:
                                run.add_picture(img_file, width=Inches(width_in * scale), height=Inches(height_in * scale))
                            
                            # Apply text wrapping and positioning properties
                            self._apply_image_wrapping(run, style)
                            
                    except Exception as e:
                        # Handle image loading errors gracefully
                        run.text = f"[Image: {os.path.basename(image_path)}]"
                else:
                    # Handle missing image
                    run.text = f"[Missing Image: {image_path}]"

    def _apply_image_wrapping(self, run, style_kwargs):
        """
        Apply text wrapping and positioning properties to an image.
        
        Args:
            run: The run containing the image
            style_kwargs: Style keyword arguments containing wrapping properties
        """
        # Get the shape (image) from the run
        if not run._element.findall('.//pic:pic', {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}):
            return  # No image found in run
        
        # Get the shape element
        shape = run._element.findall('.//wp:anchor', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
        if not shape:
            # If no anchor found, the image is inline - we need to convert it to floating
            # This is a complex operation that requires creating a new anchor element
            # For now, we'll only apply wrapping to images that are already floating
            return
        
        shape = shape[0]
        
        # Apply text wrapping
        wrap_text = style_kwargs.get("wrap_text")
        if wrap_text:
            self._set_text_wrapping(shape, wrap_text)
        
        # Apply positioning
        horizontal_align = style_kwargs.get("horizontal_align")
        vertical_align = style_kwargs.get("vertical_align")
        if horizontal_align or vertical_align:
            self._set_image_positioning(shape, horizontal_align, vertical_align)
        
        # Apply distance from text
        distance = style_kwargs.get("distance_from_text")
        if distance:
            self._set_distance_from_text(shape, distance)

    def _set_text_wrapping(self, shape, wrap_mode):
        """
        Set the text wrapping mode for an image.
        
        Args:
            shape: The shape element
            wrap_mode: The wrapping mode string
        """
        # Define the wrapping mode mappings
        wrap_map = {
            "inline": "inline",
            "square": "square",
            "tight": "tight",
            "through": "through",
            "top_and_bottom": "topAndBottom",
            "behind": "behind",
            "in_front": "inFront"
        }
        
        if wrap_mode in wrap_map:
            # Find or create the wrap element
            wrap_elem = shape.find('.//wp:wrap', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
            if wrap_elem is None:
                # Create wrap element if it doesn't exist
                from docx.oxml import parse_xml
                wrap_xml = f'<wp:wrap xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" type="{wrap_map[wrap_mode]}"/>'
                wrap_elem = parse_xml(wrap_xml)
                shape.append(wrap_elem)
            else:
                # Update existing wrap element
                wrap_elem.set('type', wrap_map[wrap_mode])

    def _set_image_positioning(self, shape, horizontal_align, vertical_align):
        """
        Set the horizontal and vertical positioning of an image.
        
        Args:
            shape: The shape element
            horizontal_align: Horizontal alignment ("left", "center", "right")
            vertical_align: Vertical alignment ("top", "middle", "bottom")
        """
        # Define alignment mappings
        horizontal_map = {
            "left": "left",
            "center": "center", 
            "right": "right"
        }
        
        vertical_map = {
            "top": "top",
            "middle": "middle",
            "bottom": "bottom"
        }
        
        # Apply horizontal alignment
        if horizontal_align and horizontal_align in horizontal_map:
            pos_h = shape.find('.//wp:positionH', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
            if pos_h is not None:
                align_elem = pos_h.find('.//wp:align', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                if align_elem is not None:
                    align_elem.text = horizontal_map[horizontal_align]
        
        # Apply vertical alignment
        if vertical_align and vertical_align in vertical_map:
            pos_v = shape.find('.//wp:positionV', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
            if pos_v is not None:
                align_elem = pos_v.find('.//wp:align', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                if align_elem is not None:
                    align_elem.text = vertical_map[vertical_align]

    def _set_distance_from_text(self, shape, distance):
        """
        Set the distance from text for an image.
        
        Args:
            shape: The shape element
            distance: Distance string (e.g., "0.1in", "10px")
        """
        # Parse the distance
        distance_in = self._parse_measurement(distance)
        if distance_in is None:
            return
        
        # Convert to EMUs (Excel Metric Units - 1 inch = 914400 EMUs)
        distance_emu = int(distance_in * 914400)
        
        # Apply to wrap margins
        wrap_elem = shape.find('.//wp:wrap', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
        if wrap_elem is not None:
            # Set all margins to the specified distance
            for margin in ['left', 'right', 'top', 'bottom']:
                margin_elem = wrap_elem.find(f'.//wp:{margin}', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                if margin_elem is not None:
                    margin_elem.text = str(distance_emu)

    def _parse_measurement(self, value):
        """
        Parse measurement strings and convert to inches.
        
        Accepts strings like '4in' or '300px' and returns inches as float.
        Supports only inches and pixels for now.
        
        Args:
            value: Measurement string (e.g., "4in", "300px") or None
            
        Returns:
            float: Measurement in inches, or None if parsing fails
        """
        if not value or not isinstance(value, str):
            return None

        value = value.strip().lower()

        if value.endswith("in"):
            try:
                return float(value.replace("in", ""))
            except ValueError:
                return None
        elif value.endswith("px"):
            try:
                px = float(value.replace("px", ""))
                return px / 96.0  # assuming 96 dpi standard
            except ValueError:
                return None
        else:
            return None 