"""
Test text block functionality
"""

import os
from docx import Document
from docxblocks import DocxBuilder
from docxblocks.constants import DEFAULT_EMPTY_VALUE_TEXT


def test_text_block(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "", "style": {"align": "center"}}
    ]
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    found = any(DEFAULT_EMPTY_VALUE_TEXT in p.text for p in doc2.paragraphs)
    assert found, "Placeholder text not found in output docx"


def test_single_newline_creates_paragraph(tmp_path):
    """Test that single newlines create new paragraphs"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "First line\nSecond line"},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]

    # Should have 2 paragraphs from the single newline
    assert len(paragraphs) == 2
    assert paragraphs[0] == "First line"
    assert paragraphs[1] == "Second line"


def test_double_newline_creates_paragraph_with_blank_line(tmp_path):
    """Test that double newlines create two separate paragraphs with blank line in between"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "First paragraph\n\nSecond paragraph"},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]

    # Should have 2 paragraphs: First, Second (blank paragraph is not counted)
    assert len(paragraphs) == 2
    assert paragraphs[0] == "First paragraph"
    assert paragraphs[1] == "Second paragraph"


def test_mixed_newlines(tmp_path):
    """Test mixed single and double newlines"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "First\nSecond\n\nThird\nFourth"},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]

    # Should have 4 paragraphs with text: First, Second, Third, Fourth
    assert len(paragraphs) == 4
    assert paragraphs[0] == "First"
    assert paragraphs[1] == "Second"
    assert paragraphs[2] == "Third"
    assert paragraphs[3] == "Fourth"





def test_newlines_create_paragraphs(tmp_path):
    """Test that all newlines create new paragraphs"""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "text", "text": "Line 1\nLine 2\nLine 3"},
    ]

    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))

    # Get all paragraphs with text content
    paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]

    # Should have 3 paragraphs from the single text block with newlines
    assert len(paragraphs) == 3
    assert paragraphs[0] == "Line 1"
    assert paragraphs[1] == "Line 2"
    assert paragraphs[2] == "Line 3" 